"""Build a Word deliverable from the thesis Markdown draft.

The converter intentionally supports the subset used by docs/tesis_apa_borrador.md:
headings, paragraphs, bold fragments, bullets, Markdown tables, images and
horizontal rules. It applies APA-friendly defaults while preserving the
document structure and figure/table captions.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE = Path("docs/tesis_apa_borrador.md")
OUT = Path("docs/deliverables/tesis_apa_tamayo.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()  # noqa: SLF001 - python-docx has no public helper
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(clean_inline(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("_", "")
    return text.strip()


def add_inline_markdown(paragraph, text: str) -> None:
    """Add text with light support for **bold** and `code` spans."""
    text = text.replace("`", "")
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if is_bold else part
        run = paragraph.add_run(content)
        run.bold = is_bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Inches(0.5)

    for name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5


def add_title_block(document: Document, lines: list[str]) -> int:
    """Render title and metadata until the first horizontal rule."""
    idx = 0
    while idx < len(lines):
        raw = lines[idx].strip()
        if raw == "---":
            document.add_page_break()
            return idx + 1
        if not raw:
            idx += 1
            continue
        if raw.startswith("> **Nota de trabajo."):
            idx += 1
            continue
        if raw.startswith("# "):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(raw[2:].strip())
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)
        else:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Inches(0)
            add_inline_markdown(p, raw.replace("  ", ""))
        idx += 1
    return idx


def is_table_start(lines: list[str], idx: int) -> bool:
    return (
        idx + 1 < len(lines)
        and lines[idx].strip().startswith("|")
        and lines[idx + 1].strip().startswith("|")
        and re.match(r"^\|?\s*:?-{3,}:?", lines[idx + 1].strip())
    )


def parse_table(lines: list[str], idx: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        line = lines[idx].strip().strip("|")
        cells = [cell.strip() for cell in line.split("|")]
        if not all(re.match(r"^:?-{3,}:?$", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, text, bold=r_idx == 0)
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
    document.add_paragraph()


def add_image(document: Document, md_line: str, base_dir: Path) -> None:
    match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", md_line.strip())
    if not match:
        return
    image_path = (base_dir / match.group(1)).resolve()
    if not image_path.exists():
        p = document.add_paragraph()
        p.add_run(f"[Imagen no encontrada: {match.group(1)}]").italic = True
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))


def build_docx() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)
    idx = add_title_block(document, lines)

    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()
        if not line:
            idx += 1
            continue
        if line == "---":
            document.add_page_break()
            idx += 1
            continue
        if is_table_start(lines, idx):
            rows, idx = parse_table(lines, idx)
            add_table(document, rows)
            continue
        if line.startswith("!["):
            add_image(document, line, SOURCE.parent)
            idx += 1
            continue
        if line.startswith("# "):
            text = clean_inline(line[2:])
            p = document.add_heading(text, level=1)
            p.paragraph_format.first_line_indent = Inches(0)
            idx += 1
            continue
        if line.startswith("## "):
            text = clean_inline(line[3:])
            p = document.add_heading(text, level=2)
            p.paragraph_format.first_line_indent = Inches(0)
            idx += 1
            continue
        if line.startswith("### "):
            text = clean_inline(line[4:])
            p = document.add_heading(text, level=3)
            p.paragraph_format.first_line_indent = Inches(0)
            idx += 1
            continue
        if line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Inches(0)
            add_inline_markdown(p, line[2:])
            idx += 1
            continue
        if re.match(r"^\d+\.\s", line):
            p = document.add_paragraph(style="List Number")
            p.paragraph_format.first_line_indent = Inches(0)
            add_inline_markdown(p, re.sub(r"^\d+\.\s", "", line))
            idx += 1
            continue
        if line.startswith("> "):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(clean_inline(line[2:]))
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            idx += 1
            continue
        p = document.add_paragraph()
        if line.startswith("**Figura") or line.startswith("**Tabla") or line.startswith("**Nota."):
            p.paragraph_format.first_line_indent = Inches(0)
        add_inline_markdown(p, line.replace("  ", ""))
        idx += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
