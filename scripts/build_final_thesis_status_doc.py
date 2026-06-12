from __future__ import annotations

import json
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor

TITLE = "Trabajo de grado — Estado final metodológico después de 15 goals"
SUBTITLE = "Modelo integrado de selección, asignación y rebalanceo de portafolios ETF"
AUTHOR = "Juan Felipe Tamayo Mejía"
OUTDIR = Path("docs/deliverables")
DOCX = OUTDIR / "trabajo_grado_estado_final_goals_1_15.docx"
PDF = OUTDIR / "trabajo_grado_estado_final_goals_1_15.pdf"
VALIDATION = OUTDIR / "trabajo_grado_estado_final_goals_1_15.validation.json"
RESULTS = Path("results/thesis_final")


def pct(x: Any) -> str:
    try:
        if pd.isna(x):
            return "—"
        return f"{float(x) * 100:.2f}%"
    except Exception:
        return "—"


def num(x: Any, digits: int = 3) -> str:
    try:
        if pd.isna(x):
            return "—"
        return f"{float(x):.{digits}f}"
    except Exception:
        return "—"


def add_p(doc: Document, text: str = "", *, style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(item)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], note: str) -> None:
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = caption.add_run(title)
    run.bold = True
    run.italic = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Nota. ").italic = True
    p.add_run(note)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 39, 97)


def cover(doc: Document) -> None:
    for text, size, bold in [
        (TITLE, 18, True),
        (SUBTITLE, 14, False),
        (AUTHOR, 12, False),
        ("Documento de síntesis generado a partir de los artefactos reproducibles del proyecto", 12, False),
        (datetime.now(timezone.utc).strftime("Generado: %Y-%m-%d %H:%M UTC"), 10, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
    doc.add_section(WD_SECTION.NEW_PAGE)


def load_tables() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    comparison = pd.read_csv(RESULTS / "tables" / "final_strategy_comparison.csv")
    intervals = pd.read_csv(RESULTS / "tables" / "final_statistical_intervals.csv")
    caps = pd.read_csv(RESULTS / "tables" / "sensitivity_cap.csv")
    rebal = pd.read_csv(RESULTS / "tables" / "sensitivity_rebalance_frequency.csv")
    return comparison, intervals, caps, rebal


def build() -> Path:
    OUTDIR.mkdir(parents=True, exist_ok=True)
    comparison, intervals, caps, rebal = load_tables()
    doc = Document()
    configure(doc)
    cover(doc)

    doc.add_heading("Resumen ejecutivo", level=1)
    add_p(
        doc,
        "Este documento consolida el estado actual del trabajo de grado después de ejecutar quince goals de desarrollo metodológico. El proyecto pasó de ser un backtest exploratorio a una arquitectura reproducible para selección, asignación y rebalanceo de portafolios ETF, con controles de disponibilidad de datos, comparación multicriterio, inferencia estadística y límites explícitos de redacción académica.",
    )
    add_p(
        doc,
        "La conclusión central es prudente: la contribución defendible es metodológica y de trazabilidad. La evidencia pública actual no permite afirmar que el modelo vence al mercado ni que la base de datos sea completamente survivor-bias-free.",
        bold_prefix="La conclusión central",
    )

    doc.add_heading("1. Resultado integrado de los 15 goals", level=1)
    goal_rows = [
        ["1–4", "Arquitectura base", "Separación entre selección, asignación y rebalanceo; ELECTRE Tri usado como clasificador."],
        ["5–7", "Backtesting y datos", "Walk-forward OOS, universo ETF público aproximado PIT y control explícito de look-ahead mediante source_available_date."],
        ["8–10", "Diagnósticos", "Etiquetas de calidad, comparación ELECTRE Tri vs FlowSort y ablations de selección/asignación."],
        ["11", "Pesos defendibles", "BWM como especificación principal, pesos manuales/equal-weight como baselines y sensibilidad aleatoria."],
        ["12–13", "Corrida final", "Comando único reproducible: python main.py --config configs/thesis_final.yaml."],
        ["14", "Inferencia", "Block bootstrap mensual, IC de CAGR/Sharpe, tests pareados, drawdown y sensibilidad por λ, pesos, cap y rebalanceo."],
        ["15", "Redacción", "Guardrails de claims para blindar la tesis ante jurado."],
    ]
    add_table(
        doc,
        "Tabla 1. Síntesis de los quince goals ejecutados",
        ["Goals", "Eje", "Resultado"],
        goal_rows,
        "Síntesis propia a partir del repositorio y de los artefactos generados en results/thesis_final/.",
    )

    doc.add_heading("2. Claims metodológicos defendibles", level=1)
    add_p(
        doc,
        "Dado que no fue posible acceder a una base institucional survivor-bias-free como CRSP, se construyó un universo ETF público aproximado point-in-time a partir de fuentes regulatorias y de mercado, incorporando fechas de disponibilidad de información y etiquetas de calidad. Esta reconstrucción no elimina por completo el riesgo de sesgo de supervivencia, pero permite reducirlo y hacerlo explícito dentro del protocolo de backtesting.",
    )
    add_bullets(
        doc,
        [
            "Se construyó un universo ETF público aproximado point-in-time.",
            "Se evitaron violaciones explícitas de look-ahead mediante source_available_date.",
            "Se etiquetó la calidad de cada observación.",
            "Se separó selección, asignación y rebalanceo.",
            "Se comparó ELECTRE Tri frente a FlowSort.",
            "Se evaluó la clasificación antes del portafolio.",
            "Se realizaron ablations para aislar fuentes de desempeño.",
        ],
    )
    add_p(
        doc,
        "No se debe afirmar que la base es completamente survivor-bias-free, que el modelo vence al mercado, que ELECTRE optimiza portafolios, que FlowSort rebalancea o que ventanas piloto de CAGR alto constituyen evidencia final.",
    )

    doc.add_heading("3. Arquitectura metodológica actual", level=1)
    add_p(
        doc,
        "La arquitectura se organiza en etapas: construcción del universo, ingeniería de criterios, clasificación multicriterio, asignación de pesos, rebalanceo y evaluación. ELECTRE Tri y FlowSort pertenecen a la capa de clasificación/ordenamiento; la asignación de capital pertenece a estrategias posteriores como EqualWeight, MinVariance o InverseVol; y el rebalanceo pertenece al motor de backtesting y política de cartera.",
    )
    add_p(
        doc,
        "Esta separación protege la tesis de atribuir performance a la etapa equivocada. Una mejora de retorno puede depender de la asignación, de la frecuencia de rebalanceo, de los costes o de la composición del universo, no necesariamente del método MCDA.",
    )

    doc.add_heading("4. Estado empírico de la corrida final", level=1)
    final = comparison[comparison["model_role"].isin(["final", "experimental"])].copy()
    rows = []
    for _, r in final.iterrows():
        rows.append([
            str(r["strategy"]),
            str(r["model_role"]),
            pct(r.get("cagr")),
            num(r.get("sharpe")),
            pct(r.get("max_drawdown")),
            str(r.get("confidence_interval", "")),
            pct(r.get("benchmark_delta")),
        ])
    add_table(
        doc,
        "Tabla 2. Resultados principales de la corrida final 2015–2025",
        ["Estrategia", "Rol", "CAGR", "Sharpe", "Max DD", "IC CAGR", "Delta vs SPY"],
        rows,
        "Los intervalos y deltas provienen de final_strategy_comparison.csv y final_statistical_intervals.csv. MaxSharpe se mantiene como experimental.",
    )
    add_p(
        doc,
        "La corrida final no sustenta una afirmación de superioridad frente a SPY. Las estrategias MCDA presentan CAGR inferior al benchmark y las notas estadísticas reportan ausencia de evidencia robusta de superioridad. Por ello, el texto de tesis debe usar formulaciones como 'no presenta evidencia robusta de superioridad' o 'presenta mejor desempeño en la muestra' solo cuando corresponda.",
    )

    doc.add_heading("5. Robustez e inferencia estadística", level=1)
    interval_rows = []
    for _, r in intervals[intervals["metric"].isin(["cagr", "sharpe", "max_drawdown"])].head(15).iterrows():
        interval_rows.append([
            str(r["strategy"]),
            str(r["metric"]),
            num(r["estimate"], 4),
            str(r["confidence_interval"]),
            str(r["statistical_note"]),
        ])
    add_table(
        doc,
        "Tabla 3. Extracto de intervalos de confianza por block bootstrap mensual",
        ["Estrategia", "Métrica", "Estimación", "Intervalo", "Nota"],
        interval_rows,
        "El bootstrap por bloques mensuales preserva dependencia temporal corta y evita basar claims solo en estimaciones puntuales.",
    )
    add_p(
        doc,
        "La inferencia incluye intervalos para CAGR, Sharpe y drawdown, tests pareados de diferencia contra SPY y sensibilidad por parámetros. Esta capa permite defender que las conclusiones se formulan con prudencia estadística.",
    )

    doc.add_heading("6. Sensibilidades y hallazgos pendientes", level=1)
    cap_rows = [[str(r.get("cap")), str(r.get("metric")), num(r.get("estimate"), 4), str(r.get("statistical_note"))] for _, r in caps.iterrows()]
    add_table(
        doc,
        "Tabla 4. Sensibilidad por cap de exposición",
        ["Cap", "Métrica", "Estimación", "Nota"],
        cap_rows,
        "Los caps 0.25, 0.35 y 0.50 fueron inviables en la corrida larga; esto se conserva como hallazgo metodológico, no se oculta.",
    )
    rebal_rows = [[str(r.get("rebalance_frequency")), str(r.get("metric")), num(r.get("estimate"), 4), pct(r.get("benchmark_delta")), str(r.get("statistical_note"))] for _, r in rebal.iterrows()]
    add_table(
        doc,
        "Tabla 5. Sensibilidad por frecuencia de rebalanceo",
        ["Frecuencia", "Métrica", "Estimación", "Delta vs SPY", "Nota"],
        rebal_rows,
        "La frecuencia anual mostró mayor CAGR que mensual/trimestral en la muestra, pero permanece por debajo de SPY y no constituye superioridad robusta.",
    )

    doc.add_heading("7. Pendientes para fortalecer el trabajo de grado", level=1)
    add_bullets(
        doc,
        [
            "Conseguir o validar una fuente institucional PIT/survivorship-aware: CRSP, Morningstar Direct, Lipper, Bloomberg, Refinitiv/LSEG o Norgate validado.",
            "Mejorar la reconstrucción pública PIT con más snapshots, cobertura de fondos liquidados/fusionados y etiquetas de calidad más granulares.",
            "Añadir criterios ETF-specific: expense ratio, AUM, spread, tracking error, beta, edad del fondo y concentración por emisor/categoría.",
            "Refinar asignación de pesos con métodos más robustos: risk parity, HRP, covariance shrinkage, caps adaptativos y volatility targeting.",
            "Profundizar la comparación ELECTRE Tri vs FlowSort: agreement, estabilidad de categorías, cambios por λ y sensibilidad de pesos.",
            "Convertir esta síntesis en capítulos finales completos: datos, metodología, diseño experimental, resultados, robustez, limitaciones, conclusiones y trabajo futuro.",
            "Preparar figuras finales: curvas de equity, drawdowns, mapa de sensibilidad, matriz de acuerdo ELECTRE/FlowSort y funnel de calidad del universo.",
        ],
    )

    doc.add_heading("8. Conclusión", level=1)
    add_p(
        doc,
        "El trabajo de grado se encuentra en estado defendible como tesis metodológica reproducible. La contribución principal es construir un protocolo auditable que integra universo público aproximado PIT, clasificación multicriterio, asignación de pesos, rebalanceo, ablations, robustez e inferencia estadística. Lo que todavía falta para elevar la tesis de metodología defendible a evidencia empírica fuerte es principalmente mejor calidad de datos PIT/survivorship-aware y criterios ETF-specific adicionales.",
    )

    doc.add_heading("Referencias operativas del repositorio", level=1)
    add_bullets(
        doc,
        [
            "configs/thesis_final.yaml",
            "results/thesis_final/run_manifest.json",
            "results/thesis_final/tables/final_strategy_comparison.csv",
            "results/thesis_final/tables/final_statistical_intervals.csv",
            "results/thesis_final/tables/final_return_difference_tests.csv",
            "docs/methodology/thesis_claims_guardrails.md",
            "docs/thesis_methodology_es.md",
        ],
    )

    doc.save(DOCX)
    return DOCX


def validate_docx(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    words = sum(len(p.split()) for p in paragraphs)
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading") and p.text.strip()]
    payload: dict[str, Any] = {
        "docx": str(path),
        "paragraphs": len(paragraphs),
        "headings": len(headings),
        "tables": len(doc.tables),
        "word_count_approx": words,
        "estimated_pages_300_words": round(words / 300, 1),
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
    }
    return payload


def convert_pdf() -> int | None:
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUTDIR), str(DOCX)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    if not PDF.exists():
        return None
    info = subprocess.run(["pdfinfo", str(PDF)], check=False, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
    for line in info.stdout.splitlines():
        if line.startswith("Pages:"):
            return int(line.split(":", 1)[1].strip())
    return None


if __name__ == "__main__":
    docx = build()
    validation = validate_docx(docx)
    try:
        validation["pdf"] = str(PDF)
        validation["exact_pdf_pages"] = convert_pdf()
    except Exception as exc:  # pragma: no cover - environment dependent
        validation["pdf_error"] = str(exc)
    VALIDATION.write_text(json.dumps(validation, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(validation, indent=2, ensure_ascii=False))
