"""Build an editable DOCX from the final enriched LaTeX thesis.

The DOCX is intended for manual academic editing. It preserves the chapter
structure, tables, figures, formulas as editable text, and an institutional
footer close to the original visual style.
"""

from __future__ import annotations

import re
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


TEX = Path(os.environ.get("TESIS_TEX", "docs/deliverables/tesis_final_tamayo_etf_electre.tex"))
ASSETS = Path("docs/deliverables/tesis_final_assets")
OUT = Path(os.environ.get("TESIS_DOCX", "docs/deliverables/tesis_final_tamayo_etf_electre_editable.docx"))
TITLE_PER_PAGE = os.environ.get("TESIS_TITLE_PER_PAGE", "0") == "1"


def clean(text: str) -> str:
    replacements = {
        r"\&": "&",
        r"\%": "%",
        r"\_": "_",
        r"--": "–",
        r"\textit": "",
        r"\textbf": "",
        r"\texttt": "",
        r"\sffamily": "",
        r"\small": "",
        r"\par": "",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\\[a-zA-Z]+(?:\[[^\]]*\])?\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"[{}]", "", text)
    text = text.replace("~", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def command_title(line: str, command: str) -> str:
    prefix = f"\\{command}{{"
    if line.startswith(prefix) and line.endswith("}"):
        return clean(line[len(prefix) : -1])
    return clean(line)


def set_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.8)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")  # noqa: SLF001
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.25)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")  # noqa: SLF001
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(192, 0, 0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def paragraph_top_border(paragraph, color: str = "C00000", size: str = "18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color)
    p_bdr.append(top)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)  # noqa: SLF001
    run._r.append(instr_text)  # noqa: SLF001
    run._r.append(fld_char2)  # noqa: SLF001


def configure_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.text = ""
        paragraph_top_border(p)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run("Escuela de Ingeniería Industrial")
        run.font.name = "Arial"
        run.font.size = Pt(9)
        p2 = footer.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.25)
        p2.paragraph_format.first_line_indent = Cm(0)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        add_page_number(p2)


def add_cover(document: Document) -> None:
    logo = ASSETS / "logo_univalle.png"
    if logo.exists():
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run().add_picture(str(logo), width=Cm(4.0))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("Programa de Ingeniería Industrial\n")
    r.bold = True
    r.font.color.rgb = RGBColor(192, 0, 0)
    p.add_run("Informe Práctica Profesional")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    title.paragraph_format.space_before = Pt(24)
    run = title.add_run("Modelo Integrado de Optimización Multicriterio para Portafolios de ETFs: Selección y\nRebalanceo")
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    for line in [
        "Presentado por:",
        "Juan Felipe Tamayo Mejía",
        "Directores del trabajo",
        "PhD. Diego Fernando Manotas Duque",
        "PhD. Orlando Joaqui Barandica",
        "Escuela de Ingeniería Industrial, Universidad del Valle",
        "Diciembre de 2025",
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(line)
        if "Juan Felipe" in line or "PhD." in line:
            run.bold = True
    document.add_page_break()


def add_text_paragraph(document: Document, text: str) -> None:
    text = clean(text)
    if not text:
        return
    p = document.add_paragraph()
    run = p.add_run(text)
    if text.startswith("Fuente:") or text.startswith("Nota:"):
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if text.startswith("Fuente:") else WD_ALIGN_PARAGRAPH.LEFT
        run.italic = True


def parse_rows(block: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw in block.splitlines():
        raw = raw.strip()
        if not raw or raw.startswith("\\") or raw in {"&", "\\toprule", "\\midrule", "\\bottomrule"}:
            continue
        if raw.startswith(("\\toprule", "\\midrule", "\\bottomrule", "\\end", "\\begin")):
            continue
        raw = raw.rstrip("\\")
        if "&" in raw:
            rows.append([clean(cell) for cell in raw.split("&")])
    return rows


def add_table(document: Document, rows: list[list[str]], caption: str | None = None) -> None:
    if caption:
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(caption)
        r.bold = True
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[j] if j < len(row) else ""
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    run.bold = i == 0
    document.add_paragraph()


def add_figure(document: Document, block: str) -> None:
    img_match = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}", block)
    cap_match = re.search(r"\\caption\{([^}]+)\}", block)
    if img_match:
        path = ASSETS / img_match.group(1)
        if path.exists():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run().add_picture(str(path), width=Inches(6.1))
    if cap_match:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(clean(cap_match.group(1)))
        r.bold = True


def extract_blocks(text: str) -> list[tuple[str, str]]:
    start = text.find(r"\section{Situación Problemática}")
    end = text.find(r"\end{document}")
    body = text[start:end]
    blocks: list[tuple[str, str]] = []
    i = 0
    while i < len(body):
        for env in ["figure", "table", "longtable", "equation"]:
            begin = rf"\begin{{{env}}}"
            if body.startswith(begin, i):
                end_marker = rf"\end{{{env}}}"
                j = body.find(end_marker, i) + len(end_marker)
                blocks.append((env, body[i:j]))
                i = j
                break
        else:
            j_candidates = [body.find(r"\begin{", i + 1)]
            j_candidates = [j for j in j_candidates if j != -1]
            j = min(j_candidates) if j_candidates else len(body)
            chunk = body[i:j]
            for line in chunk.splitlines():
                line = line.strip()
                if line:
                    blocks.append(("line", line))
            i = j
    return blocks


def build() -> None:
    tex = TEX.read_text(encoding="utf-8")
    document = Document()
    set_margins(document)
    configure_styles(document)
    configure_footer(document)
    add_cover(document)

    for kind, content in extract_blocks(tex):
        if kind == "line":
            if content.startswith(r"\section{"):
                if TITLE_PER_PAGE:
                    document.add_page_break()
                document.add_heading(command_title(content, "section"), level=1)
            elif content.startswith(r"\subsection{"):
                document.add_heading(command_title(content, "subsection"), level=2)
            elif content.startswith(r"\subsubsection{"):
                document.add_heading(command_title(content, "subsubsection"), level=3)
            elif content.startswith((r"\fuente", r"\label", r"\toprule", r"\midrule", r"\bottomrule")):
                if content.startswith(r"\fuente"):
                    add_text_paragraph(document, "Fuente: " + clean(content))
            elif content.startswith("\\"):
                continue
            else:
                add_text_paragraph(document, content)
        elif kind in {"table", "longtable"}:
            caption = None
            m = re.search(r"\\caption\{([^}]+)\}", content)
            if m:
                caption = clean(m.group(1))
            rows = parse_rows(content)
            add_table(document, rows, caption)
        elif kind == "figure":
            add_figure(document, content)
        elif kind == "equation":
            formula = content.replace(r"\begin{equation}", "").replace(r"\end{equation}", "")
            formula = re.sub(r"\\label\{[^}]+\}", "", formula).strip()
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(clean(formula))
            run.font.name = "Cambria Math"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
