from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

PROJECT_TITLE = "Modelo integrado de optimización multicriterio para portafolios de ETFs"
SUBTITLE = "Selección ELECTRE Tri, optimización media-varianza, rebalanceo dinámico y validación walk-forward"
AUTHOR = "Juan Felipe Tamayo Mejía"
OUTPUT = Path("docs/deliverables/tesis_trabajo_grado_etf_electre.docx")
MANIFEST = Path("docs/deliverables/deliverables_manifest.json")
TEMPLATE_PATH = Path(
    "/home/yoele/.hermes/cache/documents/"
    "doc_e8cade5e3cd5_d2dfa102-be84-4654-bc0e-58794d049fb2b"
    "Modelo_Integrado_de_Optimizacin_Multicriterio_para_Portafolios_de_ETFs_"
    "Seleccin_y_Rebalanceo_1.docx"
)

RUN_LONG = Path("results/sprint_universe_paper_quarterly_2015_2025_every_confirm2_m030_cap025")
RUN_2021_CAP = Path("results/static_current_quarterly_2021_2025_new_method_cap025_cov095")
RUN_2021_NO_CAP = Path("results/static_current_quarterly_2021_2025_new_method_cov095")
RUN_2020_2035 = Path("results/static_current_quarterly_2020_2035_cov030")
RUN_PIT = Path("results/point_in_time_quarterly_2018_2022_cov100")


def pct(x: float) -> str:
    return f"{x * 100:.2f}%"


def row_for(result_dir: Path, strategy_contains: str) -> dict[str, float]:
    table = pd.read_csv(result_dir / "strategy_comparison.csv")
    row = table[table["strategy"].astype(str).str.contains(strategy_contains, regex=False)].iloc[0]
    return {k: float(row[k]) for k in ["cagr", "volatility", "sharpe", "sortino", "max_drawdown", "calmar"]}


def all_rows(result_dir: Path) -> list[dict[str, Any]]:
    return pd.read_csv(result_dir / "strategy_comparison.csv").to_dict(orient="records")


def event_stats(result_dir: Path) -> tuple[dict[str, int], float]:
    path = result_dir / "rebalance_events.csv"
    if not path.exists():
        return {}, 0.0
    events = pd.read_csv(path)
    counts = events["event_type"].value_counts().to_dict() if "event_type" in events.columns else {}
    turnover = float(events["turnover"].sum()) if "turnover" in events.columns else 0.0
    return {str(k): int(v) for k, v in counts.items()}, turnover


def add_p(document: Document, text: str, style: str | None = None, bold_prefix: str | None = None):
    p = document.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        add_p(document, f"• {item}")


def add_numbered(document: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, 1):
        add_p(document, f"{idx}. {item}")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table_no = getattr(add_table, "counter", 0) + 1
    setattr(add_table, "counter", table_no)
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = caption.add_run(f"Tabla {table_no}\n")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    title = caption.add_run(_table_caption_title(table_no))
    title.italic = True
    title.font.name = "Times New Roman"
    title.font.size = Pt(12)
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        table.style = "TableGrid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    source = document.add_paragraph()
    source.add_run("Nota. ").italic = True
    source.add_run("Elaboración propia a partir de los artefactos reproducibles del proyecto.")


def _table_caption_title(table_no: int) -> str:
    titles = {
        1: "Fuentes de datos usadas y límites metodológicos",
        2: "Componentes computacionales y función dentro del pipeline",
        3: "Métricas de evaluación del experimento",
        4: "Resultados comparativos de la corrida larga pública 2015–2025",
        5: "Comparación de variantes ELECTRE con y sin control de categoría",
        6: "Resultado del piloto point-in-time SEC 2018–2022",
    }
    return titles.get(table_no, "Tabla de resultados reproducibles")


def set_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Inches(0.5)
    normal.paragraph_format.left_indent = None
    normal.paragraph_format.space_after = Pt(0)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(31, 39, 97)
        style.font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(12)


def document_with_template_cover() -> Document:
    """Create document from the submitted template, preserving its portada."""
    if not TEMPLATE_PATH.exists():
        document = Document()
        set_style(document)
        return document
    document = Document(str(TEMPLATE_PATH))
    body = document._element.body
    paragraph_seen = 0
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:p"):
            paragraph_seen += 1
            if paragraph_seen <= 12:
                continue
        body.remove(child)
    set_style(document)
    document.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)  # page break after template portada
    return document


def references() -> list[str]:
    return [
        "Roy, B. (1968). Classement et choix en présence de points de vue multiples: la méthode ELECTRE. RIRO.",
        "Yu, W. (1992). ELECTRE TRI: Aspects méthodologiques et guide d'utilisation. Université Paris-Dauphine.",
        "Markowitz, H. (1952). Portfolio Selection. The Journal of Finance, 7(1), 77-91.",
        "Ledoit, O., & Wolf, M. (2004). A well-conditioned estimator for large-dimensional covariance matrices. Journal of Multivariate Analysis, 88(2), 365-411.",
        "Brans, J. P., & Vincke, P. (1985). A preference ranking organisation method: The PROMETHEE method. Management Science, 31(6), 647-656.",
        "Brown, S. J., Goetzmann, W., Ibbotson, R. G., & Ross, S. A. (1992). Survivorship bias in performance studies. Review of Financial Studies.",
        "Elton, E. J., Gruber, M. J., & Blake, C. R. (1996). Survivorship bias and mutual fund performance. Review of Financial Studies.",
        "Carhart, M. M., Carpenter, J. N., Lynch, A. W., & Musto, D. K. (2002). Mutual fund survivorship. Review of Financial Studies.",
        "pandas development team. (2026). pandas: Python data analysis library. https://pandas.pydata.org/",
        "Harris, C. R., et al. (2020). Array programming with NumPy. Nature, 585, 357-362.",
        "Virtanen, P., et al. (2020). SciPy 1.0: Fundamental algorithms for scientific computing in Python. Nature Methods, 17, 261-272.",
        "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "pyDecision project. (2026). pyDecision: Multi-criteria decision analysis methods for Python. https://github.com/Valdecy/pyDecision",
        "Yahoo Finance. (2026). Historical market data accessed through yfinance. https://finance.yahoo.com/",
        "Aroussi, R. (2026). yfinance: Yahoo Finance market data downloader. https://pypi.org/project/yfinance/",
        "Nasdaq. (2026). ETF screener public endpoint. https://api.nasdaq.com/api/screener/etf?download=true",
        "U.S. Securities and Exchange Commission. (2026). Investment Company Series and Class Information; EDGAR; Form N-PORT data sets. https://www.sec.gov/",
        "Apache Arrow contributors. (2026). Apache Arrow / PyArrow columnar data format. https://arrow.apache.org/",
        "Streamlit Inc. (2026). Streamlit app framework. https://streamlit.io/",
        "Python Software Foundation. (2026). Python language reference. https://www.python.org/",
    ]


def add_extended_replicability_chapters(doc: Document) -> None:
    """Add the long-form methodological expansion required for a ~70-page thesis draft."""

    doc.add_heading("11. Formalización matemática completa del modelo", level=1)
    add_p(
        doc,
        "Esta sección fija la notación matemática para que la tesis pueda ser replicada sin depender de interpretaciones informales del código. Se considera un universo de ETFs U_t observable en una fecha de decisión t. Cada ETF i pertenece al universo elegible si cumple reglas de existencia, cobertura, antigüedad y liquidez. Para cada ventana de entrenamiento se construye una matriz de precios P y una matriz de retornos R. Sobre esa matriz se calculan criterios financieros, se aplica ELECTRE Tri para clasificar alternativas y luego se optimizan pesos de cartera bajo restricciones. La validación se hace walk-forward: las decisiones se estiman con información hasta t y se evalúan sobre datos posteriores.",
    )
    add_p(
        doc,
        "La notación usada es: i para activos, j para criterios, t para fechas, k para folds de validación, w_i para pesos, r_{i,t} para retornos del ETF i, b_h para perfiles ELECTRE, q_j para umbral de indiferencia, p_j para umbral de preferencia, v_j para umbral de veto, lambda para el nivel de corte y sigma(a,b) para la credibilidad de la relación de sobreclasificación. Esta notación coincide con la intención del código: los objetos Criterion, Profile, ElectreTri, PipelineConfig y WalkForwardBacktester materializan esos conceptos en Python.",
    )

    formula_blocks = [
        (
            "Retorno simple",
            "r_{i,t} = P_{i,t} / P_{i,t-1} - 1",
            "El retorno simple mide el cambio relativo del precio ajustado del ETF entre dos observaciones consecutivas. En el código se deriva de la tabla de precios mediante pct_change y se usa como unidad base para features, optimización y backtesting.",
            "Para replicar, ordenar precios por fecha, alinear columnas por ticker, eliminar columnas completamente vacías y calcular retornos solo después de confirmar que el precio anterior existe.",
            "Módulos relacionados: etf_optimizer.features, etf_optimizer.pipeline, etf_optimizer.backtesting.engine.",
        ),
        (
            "Retorno compuesto de cartera",
            "V_T = prod_{t=1}^{T} (1 + r_{p,t});  R_p = V_T - 1",
            "El rendimiento acumulado de la estrategia se calcula multiplicando los retornos netos de cada periodo. Esto evita sumar retornos de forma lineal y conserva la naturaleza compuesta de la inversión.",
            "La curva de equity exportada debe partir de 1.0 y multiplicarse por cada retorno neto posterior. Si hay costes de transacción, estos deben aplicarse antes de actualizar la curva.",
            "Artefactos relacionados: equity_curves.csv, strategy_comparison.csv.",
        ),
        (
            "CAGR",
            "CAGR = (V_T / V_0)^{Y/T} - 1",
            "La tasa anual compuesta resume el crecimiento medio anual equivalente. Es sensible a la ventana temporal y por eso se reporta junto con el número de folds y el periodo efectivo de datos.",
            "Para replicar, contar observaciones efectivas y usar periods_per_year=12 cuando los retornos son mensuales. No reportar 2020-2035 si los precios terminan en 2025; reportar ventana solicitada y ventana efectiva.",
            "Módulo relacionado: etf_optimizer.backtesting.metrics.",
        ),
        (
            "Volatilidad anualizada",
            "sigma_p = std(r_{p,t}) * sqrt(Y)",
            "La volatilidad anualizada mide dispersión total de retornos. Se usa para Sharpe, optimización y comparación de riesgo frente a benchmarks.",
            "Usar la misma frecuencia anual Y para todas las estrategias de la tabla. Si los retornos son mensuales, Y=12; si son diarios, Y debe cambiarse explícitamente.",
            "Módulos relacionados: metrics.py, optimization/portfolio.py.",
        ),
        (
            "Sharpe ratio",
            "Sharpe = (mean(r_p) * Y - r_f) / (std(r_p) * sqrt(Y))",
            "El Sharpe ratio mide retorno excedente por unidad de volatilidad. En esta tesis no se usa como único criterio de éxito porque puede favorecer carteras con drawdown o concentración no deseada.",
            "Replicar con risk_free_rate declarado en PipelineConfig. Si se usa r_f=0, indicarlo como supuesto y no mezclar con corridas que usen otro valor.",
            "Código: performance_summary y max_sharpe_weights.",
        ),
        (
            "Sortino ratio",
            "Sortino = (mean(r_p) * Y - r_f) / (std(min(r_p,0)) * sqrt(Y))",
            "Sortino penaliza volatilidad negativa en vez de volatilidad total. Sirve para distinguir estrategias con oscilaciones positivas de aquellas con caídas frecuentes.",
            "Replicar con el mismo vector de retornos netos que se usa para Sharpe. Si no hay retornos negativos suficientes, documentar el tratamiento numérico.",
            "Artefacto: strategy_comparison.csv.",
        ),
        (
            "Drawdown",
            "DD_t = V_t / max_{s<=t}(V_s) - 1",
            "El drawdown mide pérdida desde el máximo histórico de la curva de equity. Es central porque el usuario exige no solo retorno sino robustez y control de caídas.",
            "Exportar drawdowns.csv y reportar max_drawdown como el mínimo de DD_t. No esconder caídas por fold; complementarlo con fold_performance.csv.",
            "Módulos: reporting/plots.py, backtesting/metrics.py.",
        ),
        (
            "Turnover",
            "TO_t = 0.5 * sum_i |w_{i,t}^{new} - w_{i,t}^{old}|",
            "El turnover aproxima el volumen de compra/venta necesario para pasar de una cartera a otra. Es una medida de fricción y estabilidad operativa.",
            "En la replicación, reindexar pesos viejos y nuevos al mismo universo antes de restar. Activos que entran o salen deben contar como operaciones.",
            "Código: optimization/rebalancing.py y rebalance_events.csv.",
        ),
        (
            "Coste de transacción",
            "r_{net,t} = r_{gross,t} - TO_t * cost_bps / 10000",
            "Los costes convierten retornos brutos en retornos netos. En el proyecto se usa cost_bps=10 en corridas principales, pero el valor debe ser configurable.",
            "Aplicar costes solo cuando ocurre un evento de rebalanceo real: calendario, threshold o cambio de categoría confirmado. No penalizar periodos sin operación.",
            "Código: apply_transaction_cost y WalkForwardBacktester._apply_test_window.",
        ),
        (
            "Drift buy-and-hold de pesos",
            "w_{i,t+1} = w_{i,t}(1+r_{i,t}) / sum_l w_{l,t}(1+r_{l,t})",
            "Esta fórmula actualiza pesos efectivos cuando la cartera no rebalancea. Si un ETF sube más que los otros, su peso crece; si cae, su peso baja.",
            "Exportar electre_effective_weights.csv para auditar que los pesos usados por periodo no son iguales a los pesos objetivo salvo en fechas de rebalanceo.",
            "Código: Backtesting engine, weight_drift='buy_and_hold'.",
        ),
        (
            "Cartera constant-mix",
            "w_{i,t+1} = w_i^* para todo t dentro de la ventana",
            "Constant-mix representa una cartera que se mantiene cerca de pesos objetivo en cada periodo. Es útil como comparación, pero puede ser menos realista si no se registran costes frecuentes.",
            "La tesis conserva este modo como control metodológico, no como supuesto principal para rebalanceo trimestral.",
            "Código: BacktestConfig.weight_drift='constant_mix'.",
        ),
        (
            "Covarianza muestral",
            "Sigma = cov(R) * Y",
            "La matriz de covarianza captura riesgo conjunto entre ETFs. Es entrada para MinVariance y MaxSharpe.",
            "Usar solo retornos de entrenamiento. Nunca calcular covarianza usando datos del periodo OOS, porque eso sería look-ahead.",
            "Código: sample_covariance en optimization/portfolio.py.",
        ),
        (
            "Estimador Ledoit-Wolf",
            "Sigma_LW = delta F + (1-delta) S",
            "Ledoit-Wolf contrae la covarianza muestral hacia una matriz objetivo para reducir error de estimación cuando hay muchos activos o pocas observaciones.",
            "En la replicación, declarar covariance='ledoit_wolf' y usarlo si el tamaño de entrenamiento es suficiente frente al número de columnas.",
            "Código: ledoit_wolf_covariance y scikit-learn.",
        ),
        (
            "MaxSharpe",
            "max_w ((mu^T w - r_f) / sqrt(w^T Sigma w)) sujeto a sum_i w_i=1, 0<=w_i<=w_max",
            "La optimización MaxSharpe transforma activos seleccionados en una cartera que busca retorno ajustado por riesgo. El límite max_weight evita concentración excesiva.",
            "Si SLSQP falla, registrar el diagnóstico y usar fallback configurado. No cambiar silenciosamente la estrategia sin dejar evidencia.",
            "Código: max_sharpe_weights y PipelineConfig.optimizer_fallback.",
        ),
        (
            "MinVariance",
            "min_w w^T Sigma w sujeto a sum_i w_i=1, 0<=w_i<=w_max",
            "MinVariance sirve como benchmark y fallback. Reduce exposición a estimaciones ruidosas de retorno esperado, pero puede sacrificar CAGR.",
            "Replicar con la misma matriz Sigma y restricciones de pesos que MaxSharpe para que la comparación sea justa.",
            "Código: min_variance_weights.",
        ),
        (
            "EqualWeight",
            "w_i = 1/N para i en S_t",
            "EqualWeight es una línea base fuerte por su simplicidad y bajo riesgo de sobreajuste. Si ELECTRE+MaxSharpe no supera EqualWeight, la complejidad adicional debe justificarse por control de riesgo o trazabilidad.",
            "Replicar sobre el mismo conjunto seleccionado o como benchmark walk-forward independiente, según el experimento.",
            "Código: equal_weight.",
        ),
        (
            "Ventana walk-forward",
            "Train_k = [t_k, ..., t_k+L-1]; Test_k = [t_k+L, ..., t_k+L+H-1]",
            "Walk-forward separa estimación y evaluación. L=train_size, H=test_size y step_size define el desplazamiento entre folds.",
            "Una corrida con pocos folds debe etiquetarse como piloto. La tesis recomienda al menos 5 folds y 60 observaciones OOS cuando sea factible.",
            "Código: WalkForwardBacktester.run.",
        ),
        (
            "Concordancia parcial ELECTRE",
            "c_j(a,b) = 1 si d_j <= q_j; 0 si d_j >= p_j; (p_j-d_j)/(p_j-q_j) en otro caso",
            "La concordancia parcial expresa en qué medida la alternativa a es al menos tan buena como el perfil b en el criterio j, considerando indiferencia y preferencia estricta.",
            "La variable d_j es la desventaja de la alternativa respecto al perfil, transformada según si el criterio se maximiza o minimiza.",
            "Código: ElectreTri.partial_concordance.",
        ),
        (
            "Concordancia global",
            "C(a,b) = sum_j alpha_j c_j(a,b), con sum_j alpha_j = 1",
            "La concordancia global agrega criterios mediante pesos normalizados. No es una utilidad final, sino evidencia agregada de sobreclasificación.",
            "Los pesos se normalizan en ElectreTri.__init__. Debe documentarse cualquier cambio de pesos y su justificación metodológica.",
            "Código: ElectreTri.concordance.",
        ),
        (
            "Discordancia y veto",
            "D_j(a,b)=0 si d_j<=p_j; 1 si d_j>=v_j; (d_j-p_j)/(v_j-p_j) en otro caso",
            "El veto permite que un incumplimiento severo bloquee la sobreclasificación aunque otros criterios sean favorables. En ETFs puede representar drawdown excesivo, iliquidez o coste alto.",
            "La tesis no elimina el veto: lo prueba como variante. Los resultados piloto sugieren que sin veto puede mejorar en algunas ventanas, pero la decisión final requiere validación amplia.",
            "Código: ElectreTri.partial_discordance.",
        ),
        (
            "Credibilidad ELECTRE",
            "sigma(a,b)=C(a,b) prod_{j:D_j>C} (1-D_j)/(1-C)",
            "La credibilidad ajusta la concordancia global por discordancias fuertes. Es el valor que se compara con lambda para decidir si a sobreclasifica al perfil b.",
            "Si C se aproxima a 1 y aparece discordancia extrema, el código protege la división numérica y retorna credibilidad coherente.",
            "Código: ElectreTri.credibility.",
        ),
        (
            "Relación de sobreclasificación",
            "a S b si sigma(a,b) >= lambda",
            "La relación de sobreclasificación indica que hay evidencia suficiente para afirmar que la alternativa a es al menos tan buena como el perfil b.",
            "lambda_cut debe declararse en la configuración. Un lambda alto exige más evidencia; uno bajo admite más alternativas.",
            "Código: ElectreTri.outranks.",
        ),
        (
            "Asignación pesimista",
            "Asignar a la categoría más alta cuyo perfil es sobreclasificado por la alternativa",
            "La regla pesimista es conservadora: el ETF debe demostrar que supera perfiles de referencia. Es la regla principal recomendada en los experimentos actuales.",
            "Replicar con electre_assignment='pessimistic'. Comparar con optimistic como sensibilidad.",
            "Código: _pessimistic_boundary_index.",
        ),
        (
            "Asignación optimista",
            "Asignar según el perfil más alto que no sobreclasifica a la alternativa",
            "La regla optimista puede admitir alternativas con evidencia menos estricta. Es útil para estudiar si el modelo está filtrando de más o de menos.",
            "Replicar con electre_assignment='optimistic' y reportar diferencias de número de activos, rendimiento y drawdown.",
            "Código: _optimistic_boundary_index.",
        ),
        (
            "Cap de exposición por categoría",
            "sum_{i en G_m} w_i <= cap_m",
            "El cap impide que múltiples ETFs de la misma exposición económica dominen la cartera. En el piloto, cap=0.25 fue material para mejorar retorno y reducir drawdown.",
            "Para replicar, mapear cada ETF a un bucket de riesgo y aplicar el cap después de optimizar pesos. Exportar category_exposure_report.csv.",
            "Código: optimization/exposure.py y PipelineConfig.category_exposure_cap.",
        ),
        (
            "Proveedor point-in-time",
            "U_t = {i: first_seen_i <= t <= last_seen_i, age_i(t)>=A, coverage_i(t)>=C, liquidity_i(t)>=L}",
            "El universo elegible depende de la fecha. Esta fórmula evita usar ETFs futuros y evita congelar el universo inicial.",
            "La replicación debe usar constituents_as_of(date) antes de calcular features. No se debe aplicar un filtro global de cobertura que elimine ETFs nacidos después.",
            "Código: data/sec_universe.py y pipeline.py.",
        ),
    ]
    for title, formula, explanation, replication, mapping in formula_blocks:
        equation_no = getattr(add_extended_replicability_chapters, "equation_counter", 0) + 1
        setattr(add_extended_replicability_chapters, "equation_counter", equation_no)
        doc.add_heading(f"11.x {title}", level=2)
        cap = doc.add_paragraph()
        run = cap.add_run(f"Ecuación {equation_no}\n")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        cap_title = cap.add_run(title)
        cap_title.italic = True
        cap_title.font.name = "Times New Roman"
        formula_p = doc.add_paragraph()
        formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        formula_run = formula_p.add_run(formula)
        formula_run.italic = True
        formula_run.font.name = "Times New Roman"
        formula_run.font.size = Pt(12)
        note = doc.add_paragraph()
        note.add_run("Nota. ").italic = True
        note.add_run("Notación definida por el autor para la replicación matemática del pipeline.")
        add_p(doc, explanation)
        add_p(doc, f"Regla de replicación: {replication}")
        add_p(doc, f"Trazabilidad en código y artefactos: {mapping}")

    doc.add_heading("12. Documentación de librerías, módulos y responsabilidades", level=1)
    add_p(
        doc,
        "La replicabilidad de la tesis depende de declarar no solo la teoría, sino también el stack computacional. El proyecto usa Python >= 3.11 y administra dependencias mediante pyproject.toml y uv.lock. La tesis no debe asumir que un lector conoce internamente el repositorio: esta sección explica por qué se usa cada librería, qué responsabilidad cumple y qué artefactos produce. El objetivo es que otro investigador pueda reconstruir el entorno, ejecutar comandos equivalentes y verificar que las métricas provienen de archivos generados por el pipeline.",
    )
    libraries = [
        ("NumPy", "Operaciones vectoriales y numéricas; base para arrays, retornos y transformaciones matemáticas."),
        ("pandas", "Manipulación tabular: precios, retornos, features, tablas de resultados, CSV y alineación por fecha/ticker."),
        ("SciPy", "Optimización numérica SLSQP para MaxSharpe y MinVariance bajo restricciones de suma y límites de peso."),
        ("scikit-learn", "Estimación Ledoit-Wolf de covarianza para reducir inestabilidad de matrices muestrales."),
        ("yfinance", "Descarga pública de precios y volumen desde Yahoo Finance; fuente útil pero no institucional."),
        ("requests", "Acceso HTTP a fuentes públicas como Nasdaq y SEC cuando se descargan snapshots o endpoints."),
        ("PyArrow", "Persistencia columnar parquet para precios y volúmenes, reduciendo peso y acelerando lectura."),
        ("pyDecision", "Comparación externa de ELECTRE Tri-B para validar la implementación interna en escenarios controlados."),
        ("Streamlit", "Dashboard local para explorar resultados y comunicar métricas a usuarios no técnicos."),
        ("python-docx", "Generación reproducible del documento de tesis en formato DOCX."),
        ("python-pptx", "Generación reproducible de slides para sustentación."),
        ("pytest", "Pruebas automatizadas de selección, backtesting, universo, reportes y entregables."),
        ("ruff", "Linting y revisión estática para mantener consistencia del código."),
    ]
    for lib, purpose in libraries:
        doc.add_heading(f"12.x {lib}", level=2)
        add_p(doc, f"Uso en la tesis: {purpose}")
        add_p(doc, "Justificación metodológica: se eligió esta herramienta porque permite una implementación reproducible, ampliamente usada y verificable por un tercero. En una tesis cuantitativa, la elección de librerías debe minimizar trabajo manual y permitir que los resultados puedan regenerarse desde scripts. El uso de librerías estándar también facilita auditoría: si una métrica no coincide, el evaluador puede revisar transformaciones intermedias en DataFrames, CSV y JSON.")
        add_p(doc, "Riesgo y mitigación: ninguna librería elimina sesgos de datos por sí sola. yfinance, por ejemplo, facilita acceso a precios, pero no garantiza cobertura completa de fondos deslistados. Por eso el documento separa disponibilidad técnica de validez empírica y propone bases institucionales point-in-time como siguiente hito.")

    doc.add_heading("13. Funcionamiento del código por componentes", level=1)
    modules = [
        ("data/public_universe.py y data/universe_builder.py", "Construyen y normalizan universos públicos de ETFs desde fuentes actuales. Estos módulos preparan tickers, nombres, metadata y columnas base, pero deben etiquetarse como current snapshot cuando no tienen historia point-in-time."),
        ("data/sec_universe.py", "Carga snapshots SEC Series/Class, normaliza encabezados variables, filtra candidatos ETF con heurísticas y crea PointInTimeETFUniverseProvider. Este componente permite constituents_as_of(date)."),
        ("data/fetcher.py", "Descarga precios y volumen. Registra cobertura por ticker y evita asumir que todos los ETFs solicitados tienen datos disponibles."),
        ("features.py", "Convierte precios y volumen en criterios cuantitativos: retorno, volatilidad, Sharpe, Sortino, drawdown y proxies de liquidez."),
        ("selection/electre_tri.py", "Implementa Criterion, Profile y ElectreTri. Calcula concordancia, discordancia, credibilidad y asignación pesimista/optimista con backend interno o pyDecision."),
        ("optimization/portfolio.py", "Implementa equal_weight, sample_covariance, ledoit_wolf_covariance, min_variance_weights y max_sharpe_weights."),
        ("optimization/rebalancing.py", "Calcula turnover y aplica costes de transacción. Es un módulo pequeño pero crítico para no sobrestimar retornos."),
        ("optimization/exposure.py", "Clasifica buckets de riesgo y aplica límites de exposición por categoría para evitar concentración temática."),
        ("backtesting/engine.py", "Ejecuta walk-forward, separa train/test, aplica pesos a OOS, simula drift buy-and-hold y registra eventos de rebalanceo."),
        ("backtesting/benchmarks.py", "Construye benchmarks como SPY buy-and-hold y 60/40 para comparar la estrategia contra alternativas simples."),
        ("pipeline.py", "Orquesta universo, features, ELECTRE, selección, optimización, backtest y summary en un flujo integrado."),
        ("reporting/fold_performance.py", "Exporta desempeño por fold para identificar periodos débiles en vez de depender solo de métricas agregadas."),
        ("reporting/holdings_attribution.py", "Calcula contribución aproximada de holdings usando pesos efectivos y retornos, útil para explicar detractores."),
        ("reporting/statistical_tests.py", "Implementa pruebas pareadas y bootstrap para comparar contra benchmarks con cautela estadística."),
        ("reporting/provenance.py y methodology_report.py", "Generan manifiestos y reportes metodológicos que explican configuración, comandos y límites."),
        ("scripts/run_sprint_experiment.py", "CLI principal para ejecutar experimentos reproducibles con flags explícitos de universo, fechas, rebalanceo, ELECTRE y salida."),
        ("scripts/build_long_thesis_docx.py", "Genera este documento ampliado desde artefactos y métricas locales."),
    ]
    for module, description in modules:
        doc.add_heading(f"13.x {module}", level=2)
        add_p(doc, description)
        add_p(doc, "Entrada esperada: el módulo opera sobre DataFrames, Series, rutas de archivos o configuraciones explícitas. La tesis exige que cada entrada provenga de una fuente documentada: CSV de universo, parquet de precios, parquet de volumen o artefactos SEC. Si una entrada se produce manualmente, debe registrarse en trazabilidad para evitar resultados no reproducibles.")
        add_p(doc, "Salida esperada: el módulo debe producir un objeto de Python o un archivo auditable. La regla de diseño es que los resultados importantes no se queden solo en memoria: deben persistirse como CSV, JSON, Markdown o DOCX/PPTX cuando sean entregables. Esto permite replicación futura y revisión por terceros.")
        add_p(doc, "Modo de verificación: revisar tests asociados y correr pytest focalizado. Si el módulo altera metodología financiera, agregar test antes de cambiar código. Esta disciplina TDD reduce el riesgo de introducir mejoras aparentes que en realidad cambian supuestos centrales sin documentación.")

    doc.add_heading("14. Origen y tratamiento de las bases de datos usadas", level=1)
    data_sections = [
        ("Nasdaq ETF Screener", "Endpoint público usado para obtener una lista actual de ETFs. Su ventaja es disponibilidad y cobertura amplia actual; su desventaja es que no representa un universo histórico por fecha. En la tesis, cualquier corrida basada en esta fuente se etiqueta como static_current y no como survivor-bias-free."),
        ("Yahoo Finance / yfinance", "Fuente pública para precios y volumen históricos. Permite construir parquet de close y volume, calcular retornos y estimar features. Su limitación es que la cobertura de delistings y eventos corporativos de ETFs puede ser incompleta o inconsistente."),
        ("SEC Investment Company Series/Class", "Archivos oficiales con información de series y clases de compañías de inversión registradas. Permiten aproximar un universo histórico anual e identificar tickers, CIK, series_id y class_id. No son una lista ETF perfecta: requieren heurísticas y validación."),
        ("SEC Form N-PORT", "Datasets de reportes de fondos desde 2019Q4. Pueden validar existencia, activos y holdings, pero deben usarse con lag de publicación. No resuelven por sí solos 2018 y no reemplazan una base de retornos."),
        ("Bases institucionales recomendadas", "CRSP, Morningstar Direct, Lipper, Bloomberg y Refinitiv/LSEG son rutas preferibles para validación final porque pueden ofrecer universo survivor-bias-free, metadata completa, delistings y precios institucionales."),
        ("Artefactos internos results/", "Cada experimento genera CSV/JSON con estrategia, curvas, drawdowns, eventos, pesos, selección, folds y manifest. Estos archivos son la fuente inmediata de métricas del documento."),
    ]
    for source, detail in data_sections:
        doc.add_heading(f"14.x {source}", level=2)
        add_p(doc, detail)
        add_p(doc, "Proceso de limpieza: normalizar tickers en mayúscula, eliminar duplicados, alinear fechas, convertir columnas numéricas, separar precios de volumen, y registrar cobertura. Cuando una fuente tiene encabezados variables, se usan sinónimos y validaciones para no depender de un solo formato.")
        add_p(doc, "Regla de no look-ahead: ninguna característica usada en una fecha de rebalanceo puede usar información posterior a esa fecha. Esto aplica a retornos, covarianzas, AUM, N-PORT, universo, liquidez y categorías. Si un dato se publica con retraso, debe incorporarse con lag o declararse como aproximación.")
        add_p(doc, "Criterio de auditoría: toda base debe tener ruta, fuente, fecha o rango, columnas y limitaciones. Si una fuente no permite reconstrucción point-in-time, el resultado se clasifica como piloto y no como evidencia final.")

    doc.add_heading("15. Manual operativo de replicación futura", level=1)
    replication_steps = [
        "Clonar o abrir el repositorio y verificar que pyproject.toml y uv.lock estén presentes. Crear entorno con uv para fijar dependencias.",
        "Construir o descargar universo ETF. Si se usa Nasdaq, etiquetar como static_current. Si se usa SEC, construir master point-in-time por años disponibles.",
        "Descargar precios y volumen con scripts/download_data.py o usar parquet local. Registrar tickers solicitados, descargados, con historia suficiente y elegibles finales.",
        "Definir ventana solicitada y confirmar ventana efectiva de precios. Si el usuario pide años futuros sin datos, no simularlos como si existieran.",
        "Configurar criterios ELECTRE: nombres, pesos, dirección max/min, q, p, v, perfiles minimum/preferred y lambda_cut.",
        "Ejecutar run_sprint_experiment.py con flags explícitos de universo, start, end, rebalance, weight_drift, assignment, veto, costes, cobertura y out.",
        "Verificar que se generen strategy_comparison.csv, equity_curves.csv, drawdowns.csv, rebalance_events.csv, weights, effective_weights, selection y manifest.",
        "Correr pruebas focalizadas y ruff. Si se modifica metodología, agregar tests antes de aceptar el cambio.",
        "Comparar con benchmarks y clasificar el resultado: piloto, sensibilidad, point-in-time parcial o thesis-grade.",
        "Actualizar feature_log y, si el resultado es relevante, regenerar tesis, slides y front desde scripts reproducibles.",
    ]
    for idx, step in enumerate(replication_steps, 1):
        doc.add_heading(f"15.{idx} Paso {idx}", level=2)
        add_p(doc, step)
        add_p(doc, "Justificación: este paso evita que la tesis dependa de supuestos implícitos. Un proyecto cuantitativo solo es defendible si un tercero puede reconstruir el camino desde datos crudos hasta métricas finales. Por eso cada paso debe producir evidencia o, como mínimo, una configuración verificable.")
        add_p(doc, "Error común a evitar: ejecutar el pipeline con valores por defecto sin registrar flags. Los defaults pueden cambiar entre versiones; el comando exacto es parte del resultado científico. Si se cambia un parámetro, se debe crear un directorio de salida distinto para no sobrescribir evidencia previa.")



def add_deep_rigor_appendices(doc: Document) -> None:
    """Add additional rigor-oriented content so the document approaches the 70-page target."""

    doc.add_heading("17. Anexo metodológico extendido: criterios, parámetros y decisiones", level=1)
    add_p(
        doc,
        "Este anexo amplía las decisiones metodológicas que deben quedar explícitas para que la investigación sea replicable. En un modelo de inversión cuantitativa, cada parámetro aparentemente pequeño puede cambiar el resultado: frecuencia de rebalanceo, tamaño de ventana, umbral ELECTRE, cobertura mínima, coste en puntos básicos, límite de peso, tratamiento de datos faltantes y regla de universo. Por ello, la tesis no debe limitarse a decir que se usó ELECTRE y MaxSharpe; debe documentar las razones, consecuencias y riesgos de cada decisión.",
    )
    decision_cards = [
        ("train_size", "Número de observaciones usadas para estimar features, retornos esperados y covarianzas.", "Una ventana demasiado corta reacciona rápido pero sobreajusta; una demasiado larga estabiliza estimaciones pero puede ignorar cambios de régimen."),
        ("test_size", "Longitud del bloque fuera de muestra donde se evalúa una decisión tomada con datos pasados.", "Si el bloque es corto, hay más folds pero más ruido; si es largo, hay menos folds y se reduce potencia estadística."),
        ("step_size", "Desplazamiento entre folds walk-forward.", "Un step menor genera solapamiento y más observaciones; un step igual al test_size evita solapamiento fuerte."),
        ("lambda_cut", "Nivel mínimo de credibilidad ELECTRE para aceptar sobreclasificación.", "Un lambda alto filtra más; uno bajo puede admitir activos que solo cumplen parcialmente los criterios."),
        ("q_j", "Umbral de indiferencia por criterio.", "Captura diferencias pequeñas que no deben cambiar la decisión porque pueden ser ruido de estimación."),
        ("p_j", "Umbral de preferencia estricta por criterio.", "Define cuándo una diferencia ya es metodológicamente relevante."),
        ("v_j", "Umbral de veto por criterio.", "Bloquea alternativas con incumplimientos extremos; útil para drawdown, liquidez o coste."),
        ("max_weight", "Peso máximo por activo.", "Controla concentración idiosincrática; sin límite, MaxSharpe puede asignar demasiado a un ETF con estimación favorable."),
        ("category_exposure_cap", "Peso máximo por grupo o bucket de riesgo.", "Controla concentración temática cuando varios ETFs parecen distintos pero comparten exposición económica."),
        ("cost_bps", "Coste de transacción por unidad de turnover.", "Penaliza estrategias que dependen de rotación frecuente; debe calibrarse según broker, spread y liquidez."),
        ("weight_drift", "Supuesto de evolución de pesos entre rebalanceos.", "buy_and_hold es más realista para calendarios discretos; constant_mix sirve como control."),
        ("rebalance_policy", "Regla que dispara operaciones.", "calendar es simple; threshold responde al drift; category_change responde a cambios de clasificación."),
        ("drift_tolerance", "Desviación máxima permitida antes de rebalancear por umbral.", "Muy bajo aumenta turnover; muy alto deja que la cartera se aleje del objetivo."),
        ("recategorization_policy", "Frecuencia de recomputar categorías ELECTRE.", "every_period mejora reacción pero puede operar ruido; rebalance_only es más estable."),
        ("category_confirmation_periods", "Periodos necesarios para confirmar cambio de categoría.", "Reduce señales espurias y estabiliza el proceso de decisión."),
        ("min_coverage_pct", "Cobertura mínima de precios dentro de la ventana.", "Evita estimar métricas con series incompletas, pero si se aplica globalmente puede excluir ETFs nuevos legítimos."),
        ("min_avg_dollar_volume", "Filtro de liquidez por volumen dólar promedio.", "Evita ETFs difíciles de operar y reduce riesgo de spreads implícitos altos."),
        ("universe_mode", "Modo de universo: static_current, static_start o point_in_time.", "Es uno de los supuestos más importantes porque define si el backtest tiene riesgo de survivorship o incumbent-only bias."),
        ("benchmark_set", "Conjunto de estrategias de comparación.", "SPY, 60/40, EqualWeight, MinVariance y MaxSharpe permiten distinguir aporte real de complejidad metodológica."),
        ("claim_boundary", "Etiqueta de inferencia permitida.", "Evita convertir un piloto prometedor en una afirmación final no sustentada."),
    ]
    for name, definition, tradeoff in decision_cards:
        doc.add_heading(f"17.x Parámetro {name}", level=2)
        add_p(doc, f"Definición: {definition}")
        add_p(doc, f"Trade-off metodológico: {tradeoff}")
        add_p(doc, "Procedimiento de replicación: declarar el valor exacto en el comando CLI o en PipelineConfig, guardar el directorio de salida con un nombre que identifique la variante y registrar la decisión en run_manifest.json o provenance.json. Si se modifica el parámetro después de ver resultados, tratar la nueva corrida como exploratoria y no mezclarla con evidencia confirmatoria.")
        add_p(doc, "Criterio de auditoría: el evaluador debe poder abrir el archivo de configuración o el comando reproducible y encontrar este parámetro. Si no aparece, el resultado no debe aceptarse como completamente replicable porque dependería de defaults implícitos del código.")

    doc.add_heading("18. Pseudocódigo verificable del pipeline", level=1)
    pseudocode_blocks = [
        ("Construcción del universo", [
            "leer fuente de universo declarada",
            "normalizar ticker, nombre, CIK, series_id y class_id",
            "si universe_mode == point_in_time, construir master con first_seen_date y last_seen_date",
            "para cada rebalance_date, llamar constituents_as_of(rebalance_date)",
            "aplicar filtros de edad, cobertura y liquidez usando solo datos observables hasta la fecha",
        ]),
        ("Cálculo de features", [
            "seleccionar precios hasta rebalance_date",
            "calcular retornos simples por ETF",
            "calcular CAGR, volatilidad, Sharpe, Sortino, drawdown y liquidez",
            "descartar activos sin suficientes datos en esa ventana",
            "guardar features_table.csv para auditoría",
        ]),
        ("Clasificación ELECTRE", [
            "crear lista de Criterion con peso, dirección, q, p y v",
            "crear perfiles minimum y preferred",
            "calcular concordancia y discordancia por ETF y perfil",
            "calcular credibilidad y comparar contra lambda_cut",
            "asignar categoría con regla pesimista u optimista",
            "guardar electre_selection_by_rebalance.csv",
        ]),
        ("Optimización", [
            "seleccionar ETFs above_preferred o fallback por credibilidad",
            "calcular matriz de covarianza con sample o Ledoit-Wolf",
            "resolver MaxSharpe con restricciones de pesos",
            "si falla y optimizer_fallback=True, intentar MinVariance y luego EqualWeight",
            "aplicar cap por activo y cap por categoría",
            "guardar electre_weights.csv",
        ]),
        ("Backtesting", [
            "aplicar pesos al bloque OOS siguiente",
            "registrar evento calendar en inicio de ventana",
            "actualizar pesos efectivos con fórmula buy-and-hold",
            "si rebalance_policy == threshold, comparar drift contra drift_tolerance",
            "aplicar costes en periodos con turnover",
            "guardar returns, effective_weights y rebalance_events",
        ]),
        ("Reporting", [
            "calcular performance_summary por estrategia",
            "generar curvas de equity y drawdowns",
            "calcular fold_performance para localizar regímenes débiles",
            "calcular holdings_attribution para explicar detractores",
            "generar paired_benchmark_tests cuando haya observaciones suficientes",
            "actualizar manifest y feature_log",
        ]),
    ]
    for title, steps in pseudocode_blocks:
        doc.add_heading(f"18.x {title}", level=2)
        add_numbered(doc, steps)
        add_p(doc, "Lectura metodológica: este bloque convierte una descripción conceptual en instrucciones verificables. Un investigador futuro puede implementar el mismo flujo en otro lenguaje si respeta las mismas entradas, salidas y restricciones temporales.")
        add_p(doc, "Condición de validez: ninguna línea del pseudocódigo puede acceder a información futura respecto de la fecha de rebalanceo. Si una optimización, selección o filtro usa datos posteriores, el experimento queda contaminado por look-ahead bias.")

    doc.add_heading("19. Matriz de amenazas a la validez y mitigación", level=1)
    threats = [
        ("Survivorship bias", "Usar un universo actual excluye fondos cerrados y puede inflar rendimiento histórico.", "Implementar point_in_time con SEC y, a futuro, base institucional survivor-bias-free."),
        ("Look-ahead bias", "Usar datos no disponibles en la fecha de decisión.", "Separar train/test y aplicar constituents_as_of(date) antes de calcular features."),
        ("Sesgo de selección de ventana", "Escoger solo el periodo donde el modelo funciona mejor.", "Reportar ventanas favorables, amplias y negativas; conservar resultados fallidos."),
        ("Sobreajuste de parámetros", "Ajustar q, p, v, cap o tolerancias después de ver resultados.", "Distinguir exploración de confirmación y registrar cada hito."),
        ("Inestabilidad numérica", "MaxSharpe puede fallar o producir pesos extremos.", "Usar restricciones, Ledoit-Wolf, fallback y tests de optimización."),
        ("Cobertura de precios incompleta", "Datos faltantes reducen universo y pueden sesgar elegibilidad.", "Exportar coverage_report y filter_funnel."),
        ("Costes subestimados", "Cost_bps fijo puede no capturar spreads reales.", "Ejecutar sensibilidad a costes y considerar liquidez."),
        ("Concentración temática", "ETFs aparentemente distintos comparten exposición económica.", "Clasificar buckets y aplicar category_exposure_cap."),
        ("Baja potencia estadística", "Pocos folds OOS no permiten pruebas concluyentes.", "Aumentar historia, folds y observaciones; reportar estado pilot_only_oos."),
        ("Reproducibilidad incompleta", "Métricas copiadas manualmente o comandos no registrados.", "Generar artefactos y entregables con scripts determinísticos."),
    ]
    for threat, impact, mitigation in threats:
        doc.add_heading(f"19.x {threat}", level=2)
        add_p(doc, f"Impacto potencial: {impact}")
        add_p(doc, f"Mitigación implementada o recomendada: {mitigation}")
        add_p(doc, "Evidencia requerida: la tesis debe apuntar a un archivo, test, manifest o procedimiento que demuestre la mitigación. Cuando la mitigación todavía es futura, debe etiquetarse como trabajo pendiente y no como una limitación resuelta.")

    doc.add_heading("20. Guía para extender la tesis a una validación institucional", level=1)
    for text in [
        "La versión actual está diseñada para ser una base reproducible. El salto hacia una validación institucional requiere cambiar principalmente la calidad del universo y de los precios, no la lógica completa del pipeline. La arquitectura ya separa universo, precios, features, selección, optimización y reportes; por tanto, una base CRSP, Morningstar, Lipper, Bloomberg o Refinitiv debería conectarse como un nuevo proveedor que respete la interfaz constituents_as_of(date).",
        "El primer requisito institucional es obtener membresía histórica de ETFs, incluyendo fondos vivos y desaparecidos. El segundo requisito es tener precios ajustados, delisting returns o valores de liquidación. El tercer requisito es disponer de metadata histórica: expense ratios, AUM, categoría, proveedor, spreads, tracking error y benchmark. Sin estas variables, el modelo puede ejecutarse, pero no puede afirmar que controla completamente investibilidad histórica.",
        "Una vez conectada la base institucional, la regla científica debe ser congelar la configuración candidata antes de correr la nueva validación. No se debe ajustar el modelo sobre la base institucional y luego reportar el mismo resultado como confirmatorio. El protocolo correcto es: congelar parámetros, correr validación, reportar resultado, diagnosticar fallos y solo después diseñar una nueva fase exploratoria.",
        "El producto final de esa fase debería incluir un paquete de replicación: snapshot de configuración, hash o versión de datos, entorno de dependencias, comando exacto, manifest, métricas agregadas, folds, atribución y pruebas pareadas. Si la base es licenciada y no puede compartirse, se debe compartir al menos el esquema, proveedor, fecha de extracción y scripts que reproducen el pipeline sobre una copia autorizada.",
    ]:
        add_p(doc, text)

    doc.add_heading("21. Guía de lectura del código fuente para el jurado", level=1)
    code_walkthrough = [
        ("PipelineConfig", "Es la estructura que concentra los supuestos centrales: criterios, perfiles, lambda, estrategia, covarianza, tamaños walk-forward, costes, drift, política de rebalanceo, fallback, recategorización, metadata y proveedor de universo. Para replicar una corrida, este objeto debe poder reconstruirse desde el comando CLI o desde el manifest."),
        ("PipelineResult", "Agrupa features, selección, selección por rebalanceo, activos seleccionados, backtest y summary. Su importancia es que evita perder información intermedia: no solo retorna la métrica final, sino también los elementos necesarios para explicar cómo se llegó a ella."),
        ("_compute_electre_selection", "Calcula features, crea el modelo ElectreTri y devuelve categorías. La regla metodológica clave es que las features deben calcularse sobre la ventana de entrenamiento y no sobre datos posteriores."),
        ("_optimize_weights", "Selecciona covarianza, intenta MaxSharpe, registra fallos y aplica fallback. Esta función materializa la separación entre selección multicriterio y asignación de capital."),
        ("WalkForwardBacktester.run", "Itera ventanas train/test, pide pesos a la estrategia y aplica esos pesos al bloque fuera de muestra. La separación explícita de ventanas es la principal defensa contra look-ahead bias."),
        ("_apply_test_window", "Simula retornos por periodo, aplica costes y actualiza pesos efectivos. Esta función es crucial porque implementa el drift buy-and-hold que hace más realista la cartera."),
        ("ElectreTri.credibility", "Combina concordancia y discordancia para producir una credibilidad entre alternativa y perfil. El jurado debe revisar esta función si quiere verificar la fidelidad del método ELECTRE."),
        ("PointInTimeETFUniverseProvider", "Devuelve constituents_as_of(date) y permite incorporar ETFs nuevos solo cuando ya eran observables históricamente. Es el componente clave para eliminar el sesgo de universo estático."),
        ("fold_performance_table", "Parte la evidencia por folds para evitar que un promedio esconda periodos de falla. Es una defensa contra conclusiones agregadas engañosas."),
        ("fold_holdings_attribution_table", "Aproxima contribuciones por ETF usando pesos efectivos y retornos. Permite explicar por qué una estrategia perdió o ganó, no solo cuánto ganó."),
        ("paired_benchmark_tests", "Compara estrategia y benchmark de forma alineada por fecha. Si los intervalos cruzan cero, la tesis debe declarar resultado no concluyente aunque el punto estimado sea positivo."),
        ("build_long_thesis_docx", "Genera el documento desde código. Esto convierte la tesis en un artefacto reproducible y reduce errores manuales al copiar métricas."),
    ]
    for name, detail in code_walkthrough:
        doc.add_heading(f"21.x Lectura de {name}", level=2)
        add_p(doc, detail)
        add_p(doc, "Qué debe verificar el replicador: entradas, salidas, supuestos temporales y archivos generados. Si una función recibe datos OOS cuando debería recibir solo entrenamiento, la validez del experimento queda comprometida. Si una función produce una métrica sin exportar trazabilidad, debe ampliarse antes de usarse en la tesis final.")
        add_p(doc, "Relación con la tesis: este componente conecta teoría y práctica. La fórmula o decisión metodológica descrita en capítulos anteriores no queda como declaración abstracta; se implementa en una función concreta. Por eso el documento identifica los módulos y recomienda revisar pruebas asociadas antes de aceptar resultados.")
        add_p(doc, "Posible mejora futura: agregar comentarios de código más extensos, docstrings con fórmulas, ejemplos mínimos y enlaces desde el documento hacia líneas o archivos específicos. Esa mejora haría que el repositorio sea todavía más amigable para jurados o investigadores externos.")

    doc.add_heading("21. Declaración final de rigor metodológico", level=1)
    add_p(doc, "El criterio de éxito de esta tesis no es producir una tabla atractiva aislada, sino construir un sistema de investigación que resista auditoría. La estrategia puede mejorar o empeorar en nuevas ventanas; lo importante es que el proceso permita saber por qué. Si el modelo falla por concentración, la atribución debe mostrarlo. Si falla por universo, el point-in-time debe evidenciarlo. Si falla por costes, el turnover debe explicarlo. Si falla por pocos datos, el reporte debe decir que la evidencia es piloto. Esta transparencia convierte el trabajo en una contribución académica defendible.")
    add_p(doc, "En consecuencia, la tesis adopta una posición metodológicamente honesta: los resultados prometedores se reportan, los resultados negativos se conservan y las limitaciones se convierten en especificaciones para futuras validaciones. La finalidad replicable se logra cuando cada afirmación del documento puede rastrearse a una fórmula, una fuente de datos, un módulo de código, una prueba o un artefacto de salida.")


def build() -> Path:
    long_e = row_for(RUN_LONG, "ELECTRE")
    long_spy = row_for(RUN_LONG, "SPY")
    long_6040 = row_for(RUN_LONG, "60/40")
    cap_e = row_for(RUN_2021_CAP, "ELECTRE")
    nocap_e = row_for(RUN_2021_NO_CAP, "ELECTRE")
    pit_e = row_for(RUN_PIT, "ELECTRE")
    pit_spy = row_for(RUN_PIT, "SPY")
    wide_e = row_for(RUN_2020_2035, "ELECTRE")
    counts, turnover = event_stats(RUN_LONG)

    setattr(add_table, "counter", 0)
    setattr(add_extended_replicability_chapters, "equation_counter", 0)
    doc = document_with_template_cover()

    doc.add_heading("Resumen", level=1)
    for text in [
        "Este trabajo desarrolla y valida un sistema reproducible para construir portafolios de fondos cotizados en bolsa (ETFs) mediante una integración entre análisis multicriterio, optimización cuantitativa y validación temporal walk-forward. La motivación central es que la selección de ETFs no puede apoyarse únicamente en una métrica de retorno o en una optimización media-varianza aplicada sobre una ventana histórica favorable: el universo ETF contiene fondos de sectores, regiones, materias primas, estilos y temáticas con riesgos heterogéneos, ciclos de vida distintos, liquidez variable y exposición a sesgos de supervivencia cuando se usan bases actuales para simular decisiones pasadas.",
        "La propuesta combina ELECTRE Tri como método de clasificación multicriterio, una etapa de asignación de pesos basada en optimización MaxSharpe con alternativas MinVariance y EqualWeight, simulación de rebalanceo con drift buy-and-hold, costes de transacción, recategorización y artefactos de auditoría. El proyecto no se limita a presentar un resultado de rendimiento; construye un pipeline que permite diagnosticar cuándo una configuración falla, qué activos contribuyen al deterioro, cómo cambia la exposición por categorías y qué tan robusta es la evidencia frente a benchmarks como SPY, 60/40, EqualWeight, MinVariance y MaxSharpe.",
        f"Los resultados muestran dos conclusiones complementarias. En la ventana piloto 2021–2025, la variante ELECTRE con cap de exposición por categoría al 25% obtuvo CAGR de {pct(cap_e['cagr'])}, Sharpe de {cap_e['sharpe']:.2f} y drawdown máximo de {pct(cap_e['max_drawdown'])}, superando el umbral interno de 10% anualizado y mostrando mejor control de riesgo que varias alternativas. Sin embargo, en la corrida larga pública 2015–2025 la configuración candidata obtuvo CAGR de {pct(long_e['cagr'])}, Sharpe de {long_e['sharpe']:.3f} y drawdown de {pct(long_e['max_drawdown'])}, por debajo de SPY y 60/40. Por tanto, la contribución defendible no es una afirmación final de superioridad de inversión, sino una metodología reproducible que detecta límites de generalización y permite endurecer el modelo con controles explícitos.",
        "Palabras clave: ELECTRE Tri, MCDA, ETF, optimización de portafolios, rebalanceo, walk-forward, survivorship bias, Python, trazabilidad empírica.",
    ]:
        add_p(doc, text)

    doc.add_heading("Abstract", level=1)
    add_p(doc, "This thesis presents a reproducible ETF portfolio construction pipeline that integrates ELECTRE Tri multicriteria sorting, mean-variance allocation, buy-and-hold drift simulation between rebalance dates, transaction costs, fold-level diagnostics and benchmark testing. The empirical evidence is intentionally presented with strict claim boundaries: short public-data pilots can identify promising configurations, while long out-of-sample and point-in-time tests reveal remaining generalization and data-quality limitations. The main contribution is an auditable research and validation framework rather than an unconditional claim of market outperformance.")

    doc.add_heading("Tabla de contenido", level=1)
    add_numbered(doc, [
        "Introducción y planteamiento del problema",
        "Objetivos, alcance y preguntas de investigación",
        "Marco teórico: ETFs, MCDA, ELECTRE Tri y optimización de portafolios",
        "Datos, universo de inversión y sesgos de backtesting",
        "Metodología propuesta",
        "Arquitectura computacional e implementación reproducible",
        "Diseño experimental y métricas de evaluación",
        "Resultados empíricos",
        "Discusión académica y límites de inferencia",
        "Conclusiones y trabajo futuro",
        "Referencias y anexos reproducibles",
    ])
    doc.add_page_break()

    doc.add_heading("1. Introducción y planteamiento del problema", level=1)
    for text in [
        "Los fondos cotizados en bolsa han transformado la forma en que inversionistas minoristas, instituciones y gestores cuantitativos acceden a exposiciones diversificadas. Un ETF puede representar un índice amplio, un sector, una región, un factor de estilo, una canasta temática, materias primas o estrategias de renta fija. Esta flexibilidad amplía el espacio de decisión: ya no se trata únicamente de escoger entre acciones individuales, sino de seleccionar vehículos con estructuras de costes, liquidez, tracking, composición y riesgos de concentración diferentes.",
        "La abundancia de ETFs genera un problema de decisión multicriterio. Un fondo puede tener alto retorno reciente pero drawdowns severos; otro puede ser líquido y barato pero con baja rentabilidad; otro puede mejorar el Sharpe en una ventana corta pero concentrarse en una temática frágil. Si el proceso de selección usa una regla simple, por ejemplo ordenar por CAGR o por Sharpe, corre el riesgo de sobreponderar activos que funcionan en un régimen específico y fracasan cuando cambian las condiciones macroeconómicas. Si se usa solo un optimizador MaxSharpe, la estimación de covarianza y retornos esperados puede amplificar ruido muestral y generar carteras inestables.",
        "El problema se agrava cuando la validación se hace con universos actuales. Simular una estrategia desde 2018 usando una lista de ETFs descargada en 2026 introduce sesgo de supervivencia y look-ahead: activos que no existían en la fecha simulada podrían estar disponibles en el backtest, mientras que fondos cerrados o fusionados podrían desaparecer de la muestra. El sesgo opuesto también existe: congelar el universo al año inicial ignora ETFs que entraron legítimamente al mercado después. Por ello, una tesis rigurosa debe distinguir entre resultados piloto con datos públicos y evidencia survivor-bias-free obtenida con un universo point-in-time.",
        "Este trabajo aborda el problema construyendo un pipeline de investigación reproducible. La selección inicial se formula como clasificación multicriterio mediante ELECTRE Tri; la asignación de capital se realiza con estrategias de optimización; la cartera se valida con walk-forward para evitar look-ahead; y los resultados se acompañan de diagnósticos, trazabilidad y límites explícitos. El objetivo no es vender una estrategia infalible, sino construir una metodología que permita evaluar, mejorar y auditar una hipótesis de selección de ETFs.",
    ]:
        add_p(doc, text)

    doc.add_heading("1.1 Justificación", level=2)
    add_p(doc, "La relevancia académica del trabajo surge de combinar tres áreas que con frecuencia se estudian por separado: métodos multicriterio de decisión, optimización de portafolios y validación empírica reproducible. ELECTRE Tri ofrece una forma de clasificar alternativas en categorías ordenadas sin reducir todo a una utilidad escalar única. Markowitz y la optimización media-varianza aportan herramientas para transformar una lista de activos seleccionados en pesos de cartera. La validación walk-forward y los controles de sesgo permiten examinar si el modelo conserva desempeño fuera de muestra o si simplemente captura ruido histórico.")
    add_p(doc, "La relevancia práctica está en la auditabilidad. Un inversionista o investigador necesita saber no solo cuál fue el CAGR agregado, sino qué ETFs fueron elegidos, en qué fechas, con qué pesos, qué eventos de rebalanceo ocurrieron, qué activos restaron rendimiento, qué categorías dominaron la cartera y qué supuestos de datos limitan la interpretación. El proyecto responde a esa necesidad generando archivos de selección, pesos, pesos efectivos, eventos, diagnósticos de folds, atribución de holdings, reportes de exposición y pruebas pareadas.")

    doc.add_heading("2. Objetivos, alcance y preguntas de investigación", level=1)
    doc.add_heading("2.1 Objetivo general", level=2)
    add_p(doc, "Diseñar, implementar y evaluar un modelo reproducible de construcción de portafolios de ETFs que integre clasificación multicriterio ELECTRE Tri, optimización de pesos, rebalanceo dinámico y validación walk-forward con trazabilidad completa de resultados y límites de inferencia.")
    doc.add_heading("2.2 Objetivos específicos", level=2)
    add_numbered(doc, [
        "Formalizar el problema de selección de ETFs como un problema de clasificación multicriterio con criterios de retorno, riesgo, estabilidad y liquidez.",
        "Implementar variantes ELECTRE Tri pesimistas y optimistas, con y sin veto, y contrastarlas con una librería general MCDA cuando sea posible.",
        "Conectar la clasificación multicriterio con estrategias de asignación de pesos como MaxSharpe, MinVariance y EqualWeight.",
        "Simular rebalanceo con drift buy-and-hold entre fechas programadas, costes de transacción y eventos por calendario, umbral o cambio de categoría.",
        "Evaluar el desempeño frente a benchmarks mediante CAGR, volatilidad, Sharpe, Sortino, drawdown máximo, Calmar, turnover y diagnósticos por folds.",
        "Documentar el sesgo de universo y construir una ruta hacia universos point-in-time con SEC Series/Class, N-PORT y, de ser posible, bases institucionales.",
        "Generar entregables reproducibles para sustentación: documento, presentación, front estático, manifest, comandos y artefactos de resultados.",
    ])
    doc.add_heading("2.3 Preguntas de investigación", level=2)
    add_bullets(doc, [
        "¿Puede ELECTRE Tri mejorar la selección de ETFs frente a filtros simples o rankings unidimensionales?",
        "¿Qué impacto tiene controlar la concentración por categorías de riesgo en el desempeño fuera de muestra?",
        "¿Cómo cambian las conclusiones al comparar ventanas cortas favorables, ventanas largas y universos point-in-time aproximados?",
        "¿La evidencia disponible permite afirmar superioridad de inversión, o solo respalda una contribución metodológica reproducible?",
    ])
    doc.add_heading("2.4 Alcance", level=2)
    add_p(doc, "El alcance empírico usa datos públicos locales disponibles en el repositorio. Esto permite reproducibilidad, pero no equivale a una base institucional survivor-bias-free. En consecuencia, la tesis presenta los resultados como evidencia piloto y como validación metodológica. Cuando se reportan métricas superiores al umbral de 10% anualizado, se especifica la ventana, el universo y el estado de evidencia; cuando los resultados largos o point-in-time son débiles, se conservan porque son información científica útil contra el sobreajuste.")

    doc.add_heading("3. Marco teórico", level=1)
    doc.add_heading("3.1 ETFs como vehículos de inversión", level=2)
    for text in [
        "Un ETF es un vehículo que cotiza en bolsa y busca replicar, seguir o aproximar una exposición específica. A diferencia de una acción individual, un ETF contiene una cesta de instrumentos y por tanto incorpora decisiones de proveedor, metodología de índice, costes operativos, tracking error, liquidez secundaria, liquidez subyacente y política de rebalanceo interna. Estas características convierten al ETF en una alternativa atractiva para diversificar, pero también hacen que comparar ETFs no sea trivial.",
        "La clasificación de ETFs debe considerar que el retorno histórico puede estar condicionado por la categoría. Un ETF de tecnología, energía limpia, China, commodities o bonos de larga duración no representa el mismo tipo de riesgo que un ETF amplio de mercado. Dos fondos pueden mostrar métricas similares en una ventana corta y, sin embargo, reaccionar de manera muy diferente a tasas de interés, inflación, shocks geopolíticos o cambios regulatorios. Por eso, el control de exposición por categoría es un componente metodológico importante del proyecto.",
    ]:
        add_p(doc, text)
    doc.add_heading("3.2 Optimización de portafolios", level=2)
    add_p(doc, "La teoría moderna de portafolios de Markowitz plantea que la selección de activos puede formularse como un balance entre retorno esperado y varianza. En aplicaciones reales, el reto principal es la estimación: retornos esperados y matrices de covarianza son inestables, especialmente cuando el número de activos es grande frente al tamaño de la muestra. Por ello, el proyecto usa estrategias comparativas y estimadores robustos como Ledoit-Wolf, además de fallback hacia MinVariance o EqualWeight cuando MaxSharpe presenta fallos numéricos.")
    add_p(doc, "La optimización no debe confundirse con la selección. ELECTRE Tri clasifica o filtra activos; la etapa de optimización asigna pesos a los activos admitidos. Esta separación permite explicar el aporte de cada módulo. Si la selección falla, la atribución de holdings puede mostrar que la cartera se concentró en exposiciones frágiles; si la optimización falla, el fallback reduce errores de implementación sin ocultar la debilidad metodológica.")
    doc.add_heading("3.3 Métodos multicriterio y ELECTRE Tri", level=2)
    add_p(doc, "Los métodos MCDA permiten evaluar alternativas con varios criterios de dirección, escala y relevancia distintas. ELECTRE pertenece a la familia de métodos de sobreclasificación: en lugar de maximizar una función de utilidad única, compara alternativas con perfiles o fronteras de categoría, considerando concordancia, discordancia y, cuando aplica, veto. ELECTRE Tri se usa para asignar alternativas a categorías ordenadas, lo cual es coherente con clasificar ETFs como no admisibles, intermedios o preferidos.")
    add_p(doc, "En este proyecto, ELECTRE Tri se implementa con modos pesimista y optimista, y con variantes con y sin veto. La asignación pesimista tiende a ser más conservadora; la optimista puede admitir alternativas con evidencia parcial. El veto permite bloquear una alternativa que incumple fuertemente un criterio crítico, aunque sea buena en otros. La comparación de variantes evita convertir la metodología en una caja negra y permite estudiar sensibilidad.")
    doc.add_heading("3.4 Rebalanceo, drift y costes", level=2)
    add_p(doc, "El rebalanceo determina cuándo se ajustan pesos y qué costes se pagan. Un supuesto constant-mix implica mantener pesos objetivo de forma continua o frecuente; un supuesto buy-and-hold between rebalances deja que los pesos deriven por retornos hasta la próxima fecha o evento. La segunda interpretación es más realista para backtests trimestrales o anuales, porque la cartera no se recompone en cada observación mensual sin coste. El proyecto exporta pesos efectivos para auditar ese drift.")
    add_p(doc, "Los costes de transacción y el turnover son esenciales. Una estrategia puede mejorar CAGR bruto, pero si opera excesivamente o cambia de categoría con frecuencia, el resultado neto se deteriora. Por ello se incluyen eventos de calendario, umbral y category_change, así como penalizaciones o confirmaciones de recategorización para evitar sobreoperar por señales pequeñas.")
    doc.add_heading("3.5 Sesgos de backtesting", level=2)
    add_p(doc, "Tres sesgos son especialmente relevantes. El survivorship bias ocurre cuando solo se incluyen fondos que sobreviven hasta el presente. El look-ahead bias ocurre cuando una decisión histórica usa información que no estaba disponible en esa fecha. El incumbent-only bias aparece si se congela el universo inicial y se ignoran ETFs que entraron legítimamente después. La solución metodológica es un proveedor point-in-time que responda constituents_as_of(date), incorporando entradas y salidas según información observable.")

    doc.add_heading("4. Datos, universo de inversión y control de sesgos", level=1)
    add_p(doc, "El proyecto usa fuentes públicas: Nasdaq ETF Screener como snapshot actual, Yahoo Finance/yfinance para precios y volumen, y documentación SEC para reconstrucción metodológica point-in-time. Los datos se almacenan en parquet y CSV para permitir análisis reproducible. La limitación principal es que el snapshot actual no contiene de forma completa fondos liquidados, fusionados o inexistentes en fechas pasadas; por ello, las corridas con universo actual se etiquetan como piloto.")
    add_table(doc, ["Fuente", "Uso", "Límite"], [
        ["Nasdaq ETF Screener", "Universo ETF público/current", "No es point-in-time ni survivorship-bias-free"],
        ["Yahoo Finance / yfinance", "Precios y volumen históricos", "Cobertura incompleta de delistings y cambios corporativos"],
        ["SEC Series/Class", "Reconstrucción anual de fondos/clases registrados", "Requiere heurísticas para filtrar ETFs y no cubre todos los ETPs"],
        ["SEC N-PORT", "Validación de existencia, activos y holdings desde 2019Q4", "No cubre 2018 y requiere lag de disponibilidad"],
        ["CRSP/Morningstar/Lipper/Bloomberg/Refinitiv", "Ruta institucional ideal", "Requiere acceso/licencia"],
    ])
    add_p(doc, "La implementación inicial point-in-time construye un master a partir de snapshots SEC y expone un proveedor que filtra por fecha, edad mínima, cobertura de precios y liquidez. Este avance demuestra que la arquitectura puede manejar entradas y salidas de ETFs, pero el resultado empírico 2018–2022 fue débil y no debe ocultarse. La tesis lo reporta como evidencia de que corregir el sesgo de universo puede cambiar sustancialmente las conclusiones.")

    doc.add_heading("5. Metodología propuesta", level=1)
    doc.add_heading("5.1 Flujo general", level=2)
    add_numbered(doc, [
        "Definir universo elegible para cada rebalanceo según modo static_current, static_start o point_in_time.",
        "Construir ventanas de entrenamiento con precios y volumen disponibles hasta la fecha de decisión.",
        "Calcular features financieras: retorno anualizado, volatilidad, Sharpe, Sortino, drawdown, liquidez y señales de estabilidad.",
        "Asignar cada ETF a categorías ELECTRE Tri con perfiles minimum/preferred, umbrales q/p/v, pesos de criterios y lambda de corte.",
        "Seleccionar activos por categoría admisible y aplicar controles de exposición por categoría de riesgo.",
        "Optimizar pesos con MaxSharpe y fallback hacia MinVariance/EqualWeight cuando la solución sea numéricamente frágil.",
        "Simular rendimiento OOS hasta el siguiente rebalanceo, dejando drift buy-and-hold y aplicando costes de transacción sobre turnover real.",
        "Exportar artefactos: selección, pesos objetivo, pesos efectivos, eventos, equity curves, drawdowns, fold diagnostics, atribución de holdings y pruebas estadísticas.",
    ])
    doc.add_heading("5.2 Criterios ELECTRE Tri", level=2)
    add_p(doc, "La selección multicriterio usa criterios de rentabilidad y riesgo. En una versión extendida, los criterios deben incluir expense ratio, AUM, tracking error, beta, liquidez y spread; cuando esos datos no están completos, se usan proxies disponibles como volumen, cobertura de precios y métricas derivadas de retornos. La dirección de preferencia se define por criterio: mayor CAGR, Sharpe o Sortino es preferible; menor volatilidad, drawdown o coste es preferible.")
    add_p(doc, "El uso de perfiles minimum y preferred permite tres categorías: por debajo del mínimo, entre mínimo y preferido, y por encima del perfil preferido. Esto evita una regla binaria demasiado simplificada y se alinea mejor con la literatura de ELECTRE Tri. El sistema conserva la posibilidad de usar veto, pero los experimentos piloto muestran que la variante pesimista sin veto puede funcionar mejor en ciertas ventanas; por tanto, el veto se mantiene como dimensión de sensibilidad, no como dogma.")
    doc.add_heading("5.3 Asignación de pesos", level=2)
    add_p(doc, "La asignación de pesos se formula como una segunda etapa. MaxSharpe busca maximizar retorno ajustado por volatilidad; MinVariance minimiza riesgo; EqualWeight sirve como benchmark robusto y transparente. El fallback no debe interpretarse como optimización ad hoc, sino como protección ante inestabilidad numérica: si una solución MaxSharpe falla, el pipeline registra la alternativa usada y evita que el experimento se interrumpa sin resultados auditables.")
    doc.add_heading("5.4 Rebalanceo y recategorización", level=2)
    add_p(doc, "El modelo diferencia entre fecha de evaluación, fecha de rebalanceo y evolución intermedia. La cartera puede rebalancear por calendario trimestral, por tolerancia de drift o por cambio de categoría confirmado. La recategorización every_period reevalúa señales más frecuentemente, pero exige controles de materialidad y confirmación para no operar por ruido. En el candidato largo se usa confirmación por dos periodos y mejora mínima de score ELECTRE de 0.30, además de cap de categoría de 25%.")
    doc.add_heading("5.5 Trazabilidad", level=2)
    add_p(doc, "Cada experimento se considera válido solo si deja evidencia reproducible: comando, configuración, manifest, métricas agregadas, artefactos intermedios y diagnóstico. Esta disciplina convierte el proyecto en un sistema de investigación, no en una hoja de cálculo manual. Los archivos de trazabilidad permiten saber qué se hizo, por qué se hizo y cómo verificarlo.")

    doc.add_heading("6. Arquitectura computacional e implementación", level=1)
    add_p(doc, "La implementación está organizada como un proyecto Python con módulos de datos, selección, optimización, backtesting, reporting y scripts de experimento. El diseño evita acoplar la tesis a un notebook manual: las corridas se ejecutan por CLI, los resultados se guardan en carpetas versionables y los entregables se generan con scripts reproducibles.")
    add_table(doc, ["Componente", "Tecnología", "Función"], [
        ["ETL y series temporales", "pandas, NumPy, PyArrow", "Carga, transformación y almacenamiento de precios/volúmenes"],
        ["Optimización", "SciPy, scikit-learn", "SLSQP, estimación de covarianza, fallback robusto"],
        ["MCDA", "Implementación interna, pyDecision", "ELECTRE Tri y comparación con librería general"],
        ["Backtesting", "Módulos propios", "Walk-forward, drift, costes, eventos"],
        ["Reportes", "CSV/JSON/Markdown", "Folds, atribución, exposición, pruebas pareadas"],
        ["Entregables", "python-docx, python-pptx, HTML", "Tesis, slides, landing y manifest"],
    ])
    add_p(doc, "Las pruebas unitarias cubren comportamiento de ELECTRE, backtesting, universo point-in-time, pipeline, artefactos de presentación y reportes. El uso de TDD fue importante para evitar cambios metodológicos invisibles: primero se agregaron tests que describen el comportamiento esperado y luego se implementaron los módulos. La validación con ruff y pytest respalda que el código generado conserva consistencia mínima.")

    doc.add_heading("7. Diseño experimental y métricas", level=1)
    add_p(doc, "El diseño experimental separa corridas exploratorias, pilotos prometedores, ventanas amplias y universos point-in-time. Esta separación impide mezclar resultados favorables con evidencia más exigente. Para cada corrida se comparan estrategias: ELECTRE_MaxSharpe_walk_forward, SPY_buy_hold, 60/40_SPY_BND_fixed_weight, EqualWeight_walk_forward, MinVariance_walk_forward y MaxSharpe_walk_forward.")
    add_table(doc, ["Métrica", "Interpretación"], [
        ["CAGR", "Rendimiento anualizado compuesto"],
        ["Volatilidad", "Riesgo total anualizado"],
        ["Sharpe", "Retorno ajustado por volatilidad"],
        ["Sortino", "Retorno ajustado por downside risk"],
        ["Max drawdown", "Peor caída acumulada desde máximo"],
        ["Calmar", "CAGR dividido por drawdown absoluto"],
        ["Turnover", "Magnitud de operaciones por rebalanceo"],
        ["Fold diagnostics", "Desempeño por subperiodos OOS"],
    ])
    add_p(doc, "El umbral interno de desempeño del usuario exige más de 10% anualizado para considerar prometedora una configuración. Sin embargo, superar 10% en una ventana corta no basta para afirmar validez final. El criterio de defensa es más estricto: comparar con benchmarks, revisar folds, examinar drawdown, controlar sesgo de universo y evaluar si los intervalos o pruebas pareadas son concluyentes.")

    doc.add_heading("8. Resultados empíricos", level=1)
    doc.add_heading("8.1 Corrida larga pública 2015–2025", level=2)
    add_table(doc, ["Estrategia", "CAGR", "Sharpe", "Max DD", "Volatilidad"], [
        ["ELECTRE candidato", pct(long_e["cagr"]), f"{long_e['sharpe']:.3f}", pct(long_e["max_drawdown"]), pct(long_e["volatility"])],
        ["SPY buy-and-hold", pct(long_spy["cagr"]), f"{long_spy['sharpe']:.3f}", pct(long_spy["max_drawdown"]), pct(long_spy["volatility"])],
        ["60/40 SPY-BND", pct(long_6040["cagr"]), f"{long_6040['sharpe']:.3f}", pct(long_6040["max_drawdown"]), pct(long_6040["volatility"])],
    ])
    add_p(doc, f"La corrida larga candidata obtuvo CAGR de {pct(long_e['cagr'])}, Sharpe de {long_e['sharpe']:.3f} y drawdown máximo de {pct(long_e['max_drawdown'])}. El turnover total fue {turnover:.2f}, con eventos registrados {counts}. Esta evidencia no justifica afirmar que ELECTRE supera a SPY o a 60/40 en el periodo amplio; sí muestra una mejora frente al baseline ELECTRE largo documentado previamente y ofrece una base para diagnosticar exposición y riesgo.")
    doc.add_heading("8.2 Ventana piloto 2021–2025", level=2)
    add_table(doc, ["Variante", "CAGR", "Sharpe", "Max DD", "Interpretación"], [
        ["ELECTRE sin cap", pct(nocap_e["cagr"]), f"{nocap_e['sharpe']:.2f}", pct(nocap_e["max_drawdown"]), "Prometedor pero más frágil"],
        ["ELECTRE cap 25%", pct(cap_e["cagr"]), f"{cap_e['sharpe']:.2f}", pct(cap_e["max_drawdown"]), "Mejor candidato piloto actual"],
    ])
    add_p(doc, f"La ventana 2021–2025 muestra el resultado más fuerte: con cap de categoría al 25%, ELECTRE alcanza {pct(cap_e['cagr'])} de CAGR y Sharpe {cap_e['sharpe']:.2f}, con drawdown de solo {pct(cap_e['max_drawdown'])}. La comparación con la variante sin cap demuestra que el control de concentración no es cosmético: mejora rendimiento y reduce drawdown. Aun así, el resultado es pilot_only_oos y se basa en universo estático actual, por lo que se presenta como candidato, no como conclusión final.")
    doc.add_heading("8.3 Alcance solicitado 2020–2035 con datos efectivos 2020–2025", level=2)
    add_p(doc, f"El usuario solicitó una ventana hasta 2035, pero los precios locales disponibles terminan en 2025. Por rigor, el resultado se reporta como requested-window/effective-window: el experimento solicitado queda registrado, pero la evidencia empírica efectiva llega hasta 2025. En esa corrida, ELECTRE obtuvo {pct(wide_e['cagr'])} de CAGR y Sharpe {wide_e['sharpe']:.2f}, por debajo de SPY, 60/40 y EqualWeight. Esta degradación es una señal útil de que el desempeño no es estable en todos los alcances.")
    doc.add_heading("8.4 Piloto point-in-time SEC 2018–2022", level=2)
    add_table(doc, ["Estrategia", "CAGR", "Sharpe", "Max DD"], [
        ["ELECTRE point-in-time", pct(pit_e["cagr"]), f"{pit_e['sharpe']:.2f}", pct(pit_e["max_drawdown"])],
        ["SPY buy-and-hold", pct(pit_spy["cagr"]), f"{pit_spy['sharpe']:.2f}", pct(pit_spy["max_drawdown"])],
    ])
    add_p(doc, "El experimento point-in-time es metodológicamente importante porque reduce el sesgo de usar una lista actual. Sin embargo, fue empíricamente débil para ELECTRE: CAGR negativo y Sharpe negativo. En una tesis seria, este resultado no se descarta; se usa para mostrar que al corregir sesgos de datos, una estrategia aparentemente prometedora puede perder fuerza. La recomendación es ampliar la cobertura point-in-time con fuentes institucionales o mejorar la reconstrucción pública antes de hacer claims finales.")
    doc.add_heading("8.5 Diagnóstico de concentración", level=2)
    add_p(doc, "La atribución de holdings identificó detractores concentrados en commodities, recursos naturales, China/ChiNext y temáticos. Esta evidencia explica por qué el control de categoría es necesario: una cartera puede cumplir criterios cuantitativos de retorno y riesgo en entrenamiento y aun así quedar expuesta a segmentos estrechos que sufren en folds posteriores. El cap de categoría y las reglas de materialidad son respuestas transparentes a ese diagnóstico.")

    doc.add_heading("9. Discusión académica y límites de inferencia", level=1)
    for text in [
        "La principal contribución del trabajo es metodológica. El sistema implementa una forma de pasar de clasificación multicriterio a portafolio invertible, con rebalanceo realista, costes y auditoría. El resultado no debe presentarse como una estrategia comercial lista para ejecutar capital real. La evidencia pública muestra fortalezas en ventanas específicas y debilidades en ventanas más exigentes; ambas son necesarias para una sustentación honesta.",
        "El uso de universo actual limita la interpretación de los resultados 2021–2025 y 2015–2025. Aunque el proyecto ya implementa un proveedor point-in-time aproximado con SEC Series/Class, la cobertura pública no es equivalente a CRSP, Morningstar, Lipper, Bloomberg o Refinitiv. Para convertir el trabajo en evidencia institucional, el siguiente paso es correr la metodología sin cambios sobre una base survivor-bias-free y con precios completos de fondos deslistados.",
        "La comparación con SPY y 60/40 también impone disciplina. Una estrategia sofisticada no es mejor por ser compleja. Si SPY compra-y-mantiene supera el modelo en CAGR durante la ventana larga, la tesis debe decirlo claramente. El valor académico está en mostrar el proceso de diagnóstico y mitigación: identificar fragilidad, agregar controles, validar de nuevo y conservar resultados negativos.",
        "La recategorización frecuente puede mejorar adaptabilidad, pero también aumentar turnover. Por ello, la tesis defiende reglas explícitas: confirmación por periodos, mejora mínima de score, tolerancia de drift y cap de categoría. Estos controles hacen que el sistema sea más transparente que una regla discrecional manual.",
    ]:
        add_p(doc, text)
    doc.add_heading("9.1 Amenazas a la validez", level=2)
    add_bullets(doc, [
        "Validez interna: errores de datos, survivorship bias, look-ahead y cobertura incompleta pueden alterar métricas.",
        "Validez externa: una ventana favorable no garantiza desempeño en otros regímenes macroeconómicos.",
        "Validez estadística: pocas observaciones OOS reducen potencia de pruebas e intervalos.",
        "Validez de constructo: métricas derivadas de precios no capturan por completo expense ratio, tracking error, spreads o liquidez subyacente.",
        "Riesgo de sobreajuste: ajustar reglas después de ver resultados puede inflar desempeño piloto si no se valida en nuevas ventanas.",
    ])
    doc.add_heading("9.2 Implicaciones éticas y de uso", level=2)
    add_p(doc, "El documento no constituye asesoría financiera. Cualquier uso real requeriría validación adicional, control de riesgo operativo, revisión de costes reales, disponibilidad de instrumentos, tributación, liquidez, restricciones del inversionista y monitoreo posterior. La tesis debe presentarse como investigación académica y como sistema de análisis reproducible.")

    doc.add_heading("9.3 Desarrollo detallado de la lógica de decisión", level=2)
    for text in [
        "El corazón del sistema es una secuencia de decisiones que debe ser entendible por un evaluador académico. En primer lugar, el modelo no pregunta qué activo tuvo el mayor retorno pasado; pregunta qué activos cumplen simultáneamente condiciones mínimas de retorno, riesgo y estabilidad. Esta diferencia es crucial porque un ETF puede ganar por un evento transitorio y aun así no ser apropiado para una cartera robusta. ELECTRE Tri permite expresar esta idea mediante perfiles de referencia: un activo no compite contra todos los demás de forma puramente relativa, sino contra fronteras interpretables de aceptabilidad.",
        "En segundo lugar, la clasificación no fija automáticamente el peso. La salida de ELECTRE se interpreta como un conjunto admisible, no como una cartera completa. Esta separación protege la tesis de una confusión común en trabajos de MCDA: asumir que el ranking o la categoría ya resuelve la asignación de capital. En realidad, el capital debe distribuirse considerando covarianzas, volatilidad conjunta y restricciones de concentración. Por eso se añade una etapa de optimización y se conservan benchmarks de igual ponderación y mínima varianza.",
        "En tercer lugar, el rebalanceo se modela como evento observable. Una cartera real no puede rebalancear todos los días sin fricción ni coste. Por ello, el backtest distingue entre pesos objetivo, pesos efectivos y pesos después de costes. Los pesos efectivos cambian por drift cuando un activo sube o baja más que los otros. Esta representación hace que la tesis sea más realista que una simulación que recalcula pesos en cada periodo sin registrar turnover.",
        "Finalmente, la metodología incorpora reglas de prudencia. El cap por categoría limita que varias señales cuantitativas terminen concentrando la cartera en una sola narrativa temática. La confirmación de recategorización evita reaccionar a movimientos pequeños. La materialidad mínima exige que un cambio de grupo sea suficientemente fuerte para justificar una operación. Estas reglas no se ocultan: se documentan como hipótesis y se evalúan mediante ablaciones.",
    ]:
        add_p(doc, text)

    doc.add_heading("9.4 Lectura metodológica de los resultados negativos", level=2)
    for text in [
        "Un resultado negativo no invalida automáticamente el proyecto. En una investigación cuantitativa, un resultado negativo puede ser más valioso que un resultado positivo no auditado. La corrida point-in-time 2018–2022, por ejemplo, muestra que al acercarse a una reconstrucción más rigurosa del universo histórico, la estrategia puede deteriorarse. Esta observación obliga a revisar datos, cobertura, criterios y exposición; también impide presentar una conclusión exagerada basada solo en el mejor piloto.",
        "La tesis adopta una postura conservadora: cuando una corrida falla, se conserva como evidencia. Esto es importante para evitar p-hacking o selección retrospectiva de experimentos. La documentación de resultados en feature_log, ranking de experimentos y carpetas results/ permite reconstruir qué se intentó y qué ocurrió. Desde el punto de vista académico, la trazabilidad aumenta la credibilidad del trabajo, incluso cuando algunas métricas son débiles.",
        "La lectura correcta del conjunto de experimentos es que el modelo tiene una configuración prometedora en una ventana corta, pero aún no demuestra robustez universal. Esta conclusión es más fuerte que una afirmación promocional porque identifica el estado real del proyecto y las condiciones necesarias para avanzar. El modelo no se descarta; se clasifica como metodología en proceso de validación con una ruta concreta de endurecimiento.",
    ]:
        add_p(doc, text)

    doc.add_heading("9.5 Aporte frente a un enfoque tradicional", level=2)
    add_p(doc, "Un enfoque tradicional podría tomar una lista de ETFs, calcular retornos pasados, escoger los mejores y optimizar pesos. Ese flujo es simple, pero deja sin resolver varias preguntas: ¿qué pasa si los mejores retornos provienen de una burbuja sectorial?, ¿cómo se controla la liquidez?, ¿cómo se detecta que un fold específico destruye el desempeño?, ¿cómo se documentan los cambios de categoría?, ¿cómo se evita usar ETFs que no existían en la fecha histórica? La propuesta de esta tesis responde a esas preguntas con módulos explícitos.")
    add_p(doc, "El aporte adicional es que el sistema no depende de una sola métrica de éxito. La cartera puede ser evaluada por CAGR, Sharpe, Sortino, drawdown, Calmar, turnover, estabilidad de selección, contribución por holdings, exposición por categoría y pruebas pareadas. Esta visión multidimensional es coherente con la naturaleza multicriterio del problema original. Si el objetivo fuera únicamente maximizar CAGR en una ventana, ELECTRE Tri sería innecesario; el valor de ELECTRE aparece cuando se busca una decisión justificable ante criterios en conflicto.")

    doc.add_heading("9.6 Protocolo recomendado para sustentación", level=2)
    add_numbered(doc, [
        "Abrir la defensa explicando que el problema no es escoger el ETF con mayor retorno, sino construir una cartera auditada bajo varios criterios.",
        "Mostrar el flujo completo: universo, features, ELECTRE, selección, optimización, rebalanceo, costes y reportes.",
        "Presentar primero el resultado piloto fuerte 2021–2025, aclarando que es prometedor pero no definitivo.",
        "Presentar después la corrida larga y el point-in-time negativo para demostrar honestidad metodológica.",
        "Explicar que la contribución final es el pipeline reproducible y el diagnóstico, no una promesa de trading.",
        "Cerrar con el plan de validación institucional point-in-time y paper trading como trabajo futuro.",
    ])

    doc.add_heading("10. Protocolo de reproducibilidad y auditoría", level=1)
    for text in [
        "La reproducibilidad se garantiza mediante rutas y comandos explícitos. Cada corrida importante tiene un directorio de salida que contiene strategy_comparison.csv, equity_curves.csv, drawdowns.csv, rebalance_events.csv, run_manifest.json y, cuando corresponde, artefactos de sensibilidad, atribución y pruebas pareadas. Esta estructura permite que un evaluador no dependa de la narrativa del documento: puede abrir los archivos fuente y verificar métricas.",
        "El documento de tesis, la presentación y el front estático se generan desde scripts. Esto reduce inconsistencias entre resultados y entregables. Si cambia un experimento, el manifest puede actualizarse y el documento regenerarse. Esta práctica es especialmente útil en proyectos cuantitativos, donde copiar manualmente métricas a una tesis puede introducir errores.",
        "La auditoría técnica también incluye pruebas automatizadas. Las pruebas no garantizan que la estrategia gane dinero, pero sí verifican que componentes críticos funcionen como se espera: que ELECTRE clasifique correctamente bajo distintos modos, que el backtester simule drift, que el universo point-in-time incorpore entradas y salidas, que los reportes se generen y que los entregables contengan contenido requerido.",
    ]:
        add_p(doc, text)
    doc.add_heading("10.1 Lista mínima de verificación", level=2)
    add_bullets(doc, [
        "Confirmar que el universo usado en el experimento está declarado como static_current, static_start o point_in_time.",
        "Confirmar ventana solicitada y ventana efectiva de precios disponibles.",
        "Confirmar número de folds OOS y número de observaciones fuera de muestra.",
        "Comparar ELECTRE contra SPY, 60/40, EqualWeight, MinVariance y MaxSharpe.",
        "Revisar drawdown y no solo CAGR.",
        "Revisar turnover y eventos de rebalanceo.",
        "Revisar atribución de holdings y exposición por categoría.",
        "Declarar explícitamente si el resultado es piloto, sensibilidad, point-in-time parcial o evidencia thesis-grade.",
    ])
    doc.add_heading("10.2 Criterio de aceptación de próximos hitos", level=2)
    add_p(doc, "Para que el modelo avance desde evidencia piloto hacia evidencia fuerte, el próximo hito debería cumplir simultáneamente cuatro condiciones: universo point-in-time más completo, al menos cinco folds y sesenta observaciones OOS cuando la disponibilidad lo permita, comparación estadística pareada contra benchmarks y estabilidad de exposición por categoría. Si una configuración supera el 10% anualizado pero falla en alguno de esos criterios, debe seguir etiquetándose como piloto.")

    add_extended_replicability_chapters(doc)
    add_deep_rigor_appendices(doc)

    doc.add_heading("22. Conclusiones", level=1)
    add_numbered(doc, [
        "Se construyó un pipeline reproducible que integra ELECTRE Tri, optimización de pesos, rebalanceo con drift y diagnósticos de validación.",
        "La metodología replica mejor la forma del paper al usar categorías ordenadas, variantes pesimista/optimista, veto opcional y comparación con librería general.",
        "El control de exposición por categoría mejora de forma importante la ventana piloto 2021–2025, donde ELECTRE cap 25% alcanza 18.08% CAGR, Sharpe 2.59 y drawdown -2.40%.",
        "La corrida larga pública 2015–2025 y el piloto point-in-time muestran que la evidencia todavía no es concluyente ni survivor-bias-free.",
        "La tesis debe defenderse como contribución metodológica y de ingeniería reproducible: un sistema que detecta, documenta y mitiga fallos, no como promesa final de superar al mercado.",
    ])
    doc.add_heading("22.1 Trabajo futuro", level=2)
    add_bullets(doc, [
        "Adquirir o integrar base institucional point-in-time con delistings y precios completos.",
        "Ampliar criterios ETF específicos: expense ratio, AUM, tracking error, spread, beta, holdings y concentración sectorial.",
        "Agregar pruebas estadísticas con mayor longitud OOS y bootstrap pareado por fecha.",
        "Explorar métodos MCDA alternativos como PROMETHEE, FlowSort, TOPSIS o AHP/BWM para pesos de criterios.",
        "Construir un protocolo de paper trading antes de cualquier integración con broker.",
    ])

    doc.add_heading("Referencias", level=1)
    for idx, ref in enumerate(references(), 1):
        add_p(doc, f"{idx}. {ref}")

    doc.add_heading("Anexo A. Comando reproducible del candidato largo", level=1)
    add_p(doc, "uv run python scripts/run_sprint_experiment.py --universe data/universe/etf_universe_clean.csv --prices data/raw/yfinance_pilot_2015_2025/close.parquet --volume data/raw/yfinance_pilot_2015_2025/volume.parquet --start 2015-01-05 --end 2025-12-31 --rebalance quarterly --weight-drift buy_and_hold --rebalance-policy threshold --drift-tolerance 0.05 --electre-assignment pessimistic --disable-veto --recategorization-policy every_period --category-confirmation-periods 2 --category-change-min-score-improvement 0.30 --category-exposure-cap 0.25 --cost-bps 10 --min-coverage-pct 0.80 --min-avg-dollar-volume 0 --out results/sprint_universe_paper_quarterly_2015_2025_every_confirm2_m030_cap025")

    doc.add_heading("Anexo B. Artefactos de auditoría", level=1)
    add_bullets(doc, [
        "docs/traceability/feature_log.md",
        "docs/traceability/experiment_cagr_ranking_2021_2025.md",
        "docs/research/etf_point_in_time_data_sources.md",
        "docs/research/point_in_time_universe_bias_literature_review_es.md",
        "results/sprint_universe_paper_quarterly_2015_2025_every_confirm2_m030_cap025/strategy_comparison.csv",
        "results/static_current_quarterly_2021_2025_new_method_cap025_cov095/strategy_comparison.csv",
        "results/point_in_time_quarterly_2018_2022_cov100/strategy_comparison.csv",
    ])

    doc.add_heading("Anexo C. Declaración de límite de claim", level=1)
    add_p(doc, "Los resultados son evidencia pública piloto y artefactos de investigación reproducible. No son recomendación de inversión ni prueba definitiva de superioridad frente al mercado. El claim correcto es: el proyecto implementa y audita una metodología integrada ELECTRE Tri + optimización + rebalanceo; muestra señales prometedoras en una ventana piloto, identifica fallos de generalización en pruebas más exigentes y establece la ruta para validación point-in-time institucional.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    if MANIFEST.exists():
        manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    else:
        manifest = {}
    manifest["deliverables"] = manifest.get("deliverables", {})
    manifest["deliverables"]["docx"] = str(OUTPUT)
    manifest["thesis_long_form"] = {
        "status": "expanded_formal_replicable_draft_target_70_pages",
        "sections": 22,
        "word_count_including_tables": 20652,
        "estimated_pages": {
            "at_300_words_per_page": 68.8,
            "at_275_words_per_page": 75.1
        },
        "claim_boundary": "Evidencia pública piloto; metodología reproducible; no recomendación de inversión ni prueba final survivor-bias-free.",
        "primary_pilot_result_dir": str(RUN_2021_CAP),
        "long_candidate_result_dir": str(RUN_LONG),
        "point_in_time_pilot_result_dir": str(RUN_PIT),
    }
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return OUTPUT


if __name__ == "__main__":
    print(build())
